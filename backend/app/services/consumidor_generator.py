"""Gerador do contrato de consumidor / transporte aereo.

Reaproveita toda a infra do ContractGenerator (template timbrado, estilos, rodape,
grade de assinatura com tags do DocuSeal) e troca so' o corpo do documento.

O texto vem literal dos modelos do escritorio (CONTRATO.docx e
"CONTRATO 2025 - COM RECLAME AQUI.docx"). O que varia por caso:
partes, companhia re', juizado, prazo de pagamento e a variante da reclamacao.
Fixos por decisao do escritorio: 25% de exito, Monica como CONTRATADA,
dados bancarios, canais oficiais e os valores do milheiro.
"""

from __future__ import annotations

import re
from datetime import date, datetime

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from app.models.contrato_consumidor import (
    ContratanteConsumidorPJ,
    ContratoConsumidorRequest,
    nome_signatario,
)
from app.services.contract_generator import ContractGenerator
from app.utils.currency import formatar_valor

CONTRATADA_NOME = "MÔNICA FURTADO PINHEIRO CHAGAS"
CONTRATADA_CPF = "077.097.876-22"
CONTRATADA_OAB = "OAB/MG 121.326"
CONTRATADA_ENDERECO = (
    "Rua Antônio de Albuquerque, 271 - 5º andar – Belo Horizonte - MG – CEP: 30112-010"
)


class ConsumidorGenerator(ContractGenerator):
    """Contrato de prestacao de servicos advocaticios — acao de consumo aerea."""

    def _build_document(  # type: ignore[override]
        self,
        data: ContratoConsumidorRequest,
        signatario_roles: list[dict] | None = None,
    ) -> Document:
        doc = self._new_document_from_template()
        self._clear_body(doc)
        self._clear_headers_footers(doc)
        self._ensure_contract_styles(doc)

        self._add_title(doc)
        self._c_partes(doc, data)
        self._c1_comunicacao(doc, data)
        self._c2_objeto(doc, data)
        self._c3_preco(doc, data)
        self._c4_pagamento(doc, data)
        self._c5_reembolso(doc, data)
        self._c6_condicoes_gerais(doc, data)
        self._c7_prazo(doc)
        self._c8_foro(doc, data)
        self._c_fechamento(doc, data, signatario_roles=signatario_roles)
        self._apply_document_standard(doc)

        return doc

    # ── Formatacao (medida no CONTRATO.docx do escritorio) ────────
    #
    # Segoe UI 11pt, justificado, entrelinha 1,15, sem espaco depois do
    # paragrafo (o original separa os blocos com linha em branco) e margens
    # 0,75 / 1,25 / 3 / 2 cm. Difere do contrato de honorarios, por isso
    # esta classe sobrescreve os tres metodos de layout.

    FONTE = "Segoe UI"
    TAMANHO = Pt(11)
    ENTRELINHA = 1.15
    # Uma linha em branco de 11pt a 1,15 ≈ 12pt: mesmo respiro do original.
    ESPACO_DEPOIS = Pt(12)

    def _apply_page_setup(self, doc) -> None:
        for section in doc.sections:
            section.top_margin = Cm(0.75)
            section.bottom_margin = Cm(1.25)
            section.left_margin = Cm(3)
            section.right_margin = Cm(2)

    def _apply_base_styles(self, doc) -> None:
        for style_name in ("Normal", "Body Text", "List Paragraph"):
            if style_name not in doc.styles:
                continue
            style = doc.styles[style_name]
            style.font.name = self.FONTE
            style._element.rPr.rFonts.set(qn("w:eastAsia"), self.FONTE)
            style.font.size = self.TAMANHO
            style.font.color.rgb = RGBColor(0, 0, 0)
            style.paragraph_format.line_spacing = self.ENTRELINHA
            style.paragraph_format.space_after = self.ESPACO_DEPOIS

    def _format_paragraph(self, paragraph, assinatura: bool = False) -> None:
        is_heading = paragraph.style and paragraph.style.name.startswith("Heading")
        is_signature = self._is_signature_paragraph(paragraph.text) or assinatura

        paragraph.paragraph_format.line_spacing = self.ENTRELINHA
        paragraph.paragraph_format.space_after = (
            Pt(0) if is_signature else self.ESPACO_DEPOIS
        )
        paragraph.paragraph_format.keep_together = True
        if is_heading:
            paragraph.paragraph_format.keep_with_next = True
        if is_signature:
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER if assinatura else WD_ALIGN_PARAGRAPH.LEFT
            )
        # Corpo justificado; titulo, assinaturas e celulas da grade nao.
        # Justificar texto curto na tabela de assinatura deforma o DocuSeal.
        parent = paragraph._element.getparent()
        in_table = parent is not None and parent.tag == qn("w:tc")
        if paragraph.alignment is None and not is_signature and not in_table:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        for run in paragraph.runs:
            run.font.name = self.FONTE
            run._element.rPr.rFonts.set(qn("w:eastAsia"), self.FONTE)
            run.font.size = self.TAMANHO
            run.font.color.rgb = RGBColor(0, 0, 0)

    @staticmethod
    def _alinhado_a_esquerda(paragraph):
        """Linhas curtas (contatos, dados bancarios) nao devem ser justificadas."""
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        return paragraph

    _PAPEL_NEGRITO = re.compile(r"(CONTRATANTES|CONTRATANTE|CONTRATADA)")

    def _paragrafo_com_destaques(self, doc: Document, texto: str, nomes: list[str]):
        """Qualificacao: negrito so nos nomes das partes e em CONTRATANTE/CONTRATADA."""
        p = doc.add_paragraph()
        nomes_ok = sorted({n for n in nomes if n}, key=len, reverse=True)
        if nomes_ok:
            padrao = "(" + "|".join(re.escape(n) for n in nomes_ok) + r"|CONTRATANTES|CONTRATANTE|CONTRATADA)"
            splitter = re.compile(padrao)
        else:
            splitter = self._PAPEL_NEGRITO

        pos = 0
        for m in splitter.finditer(texto):
            if m.start() > pos:
                p.add_run(texto[pos:m.start()]).bold = False
            p.add_run(m.group(0)).bold = True
            pos = m.end()
        if pos < len(texto):
            p.add_run(texto[pos:]).bold = False
        return p

    def _grid_assinaturas(self, doc: Document, entries: list[tuple[str, list[str]]]) -> None:
        """Assinaturas 2 por linha. Rotulo e' uma lista de linhas (nome / CPF / papel).

        Proprio deste contrato: nao mexe no grid do contrato de honorarios.
        """
        if not entries:
            return

        table = doc.add_table(rows=(len(entries) + 1) // 2, cols=2)
        table.autofit = True
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for idx, (tag, linhas) in enumerate(entries):
            cell = table.cell(idx // 2, idx % 2)
            cell.paragraphs[0].text = tag  # tag do DocuSeal
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for linha in linhas:
                p = cell.add_paragraph(linha)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.add_paragraph()

    # ── Concordancia ──────────────────────────────────────────────

    @staticmethod
    def _sp(data: ContratoConsumidorRequest, singular: str, plural: str) -> str:
        """Escolhe a redacao conforme o numero de contratantes."""
        return plural if data.plural else singular

    @staticmethod
    def _genero_de(contratante) -> str:
        """PJ concorda no feminino ('a empresa contratante')."""
        if isinstance(contratante, ContratanteConsumidorPJ):
            return "F"
        return contratante.genero

    @classmethod
    def _genero_plural(cls, data: ContratoConsumidorRequest) -> str:
        """'a' quando todas sao femininas; 'o' (masculino) em qualquer outro caso."""
        return "a" if all(cls._genero_de(c) == "F" for c in data.contratantes) else "o"

    @staticmethod
    def _nacionalidade(informada: str, genero_f: bool) -> str:
        """"Brasileira (o)" e' o rotulo generico do formulario — no contrato precisa
        concordar com o genero. Nacionalidade estrangeira digitada e' preservada."""
        texto = informada.strip()
        generica = texto.lower().replace(" ", "") in {
            "",
            "brasileira(o)",
            "brasileiro(a)",
            "brasileira",
            "brasileiro",
        }
        return ("brasileira" if genero_f else "brasileiro") if generica else texto

    # ── Partes ────────────────────────────────────────────────────

    def _qualificacao_pj(self, c: ContratanteConsumidorPJ) -> str:
        rep_f = c.representante_genero == "F"
        rep_nacionalidade = self._nacionalidade(c.representante_nacionalidade, rep_f)
        partes = [
            c.razao_social.upper(),
            "pessoa jurídica de direito privado",
            f"inscrita no CNPJ sob o n.º {c.cnpj}",
            f"com sede à {c.endereco}",
        ]
        if c.email:
            partes.append(f"email: {c.email}")
        partes.append(
            f"neste ato representada por {c.representante_nome.upper()}, {rep_nacionalidade}, "
            f"{'inscrita' if rep_f else 'inscrito'} no CPF sob o n.º {c.representante_cpf}"
        )
        return ", ".join(partes)

    def _qualificacao(self, contratante, *, com_endereco: bool) -> str:
        if isinstance(contratante, ContratanteConsumidorPJ):
            # A sede sempre acompanha a empresa: nao entra no agrupamento de domicilio.
            return self._qualificacao_pj(contratante)

        genero_f = contratante.genero == "F"
        nacionalidade = self._nacionalidade(contratante.nacionalidade, genero_f)
        partes = [
            contratante.nome.upper(),
            nacionalidade,
            f"{'inscrita' if genero_f else 'inscrito'} no CPF sob o n.º {contratante.cpf}",
        ]
        if contratante.rg:
            partes.append(f"RG n.º {contratante.rg}")
        if com_endereco:
            residente = "residente e domiciliada" if genero_f else "residente e domiciliado"
            partes.append(f"{residente} à {contratante.endereco}")
            if contratante.email:
                partes.append(f"email: {contratante.email}")
            if contratante.celular:
                partes.append(f"celular: {contratante.celular}")
        return ", ".join(partes)

    def _c_partes(self, doc: Document, data: ContratoConsumidorRequest) -> None:
        # Agrupar o domicilio so' faz sentido entre pessoas fisicas: PJ tem sede.
        so_pf = all(not isinstance(c, ContratanteConsumidorPJ) for c in data.contratantes)
        enderecos = {c.endereco.strip() for c in data.contratantes}
        endereco_comum = so_pf and len(enderecos) == 1 and data.plural

        quals = [
            self._qualificacao(c, com_endereco=not endereco_comum)
            for c in data.contratantes
        ]
        texto = " e ".join(quals) if len(quals) > 1 else quals[0]

        if endereco_comum:
            g = self._genero_plural(data)
            primeiro = data.contratantes[0]
            texto += f", residentes e domiciliad{g}s à {primeiro.endereco}"
            emails = [c.email for c in data.contratantes if c.email]
            celulares = [c.celular for c in data.contratantes if c.celular]
            if emails:
                texto += f", email: {'; '.join(emails)}"
            if celulares:
                texto += f", celular: {'; '.join(celulares)}"

        if data.plural:
            g = self._genero_plural(data)
            texto += f", doravante denominad{g}s apenas CONTRATANTES;"
        else:
            g = "a" if self._genero_de(data.contratantes[0]) == "F" else "o"
            texto += f", doravante denominad{g} apenas CONTRATANTE;"

        nomes_contratante: list[str] = []
        for c in data.contratantes:
            if isinstance(c, ContratanteConsumidorPJ):
                nomes_contratante.append(c.razao_social.upper())
                nomes_contratante.append(c.representante_nome.upper())
            else:
                nomes_contratante.append(c.nome.upper())

        self._paragrafo_com_destaques(doc, f"CONTRATANTES: {texto}", nomes_contratante)
        doc.add_paragraph("E")
        self._paragrafo_com_destaques(
            doc,
            f"{CONTRATADA_NOME}, {CONTRATADA_OAB}, CPF n° {CONTRATADA_CPF}, "
            f"com escritório na {CONTRATADA_ENDERECO}, doravante denominada apenas CONTRATADA.",
            [CONTRATADA_NOME],
        )
        doc.add_paragraph(
            "As partes acima qualificadas têm entre si contratado o seguinte a Contrato "
            "de Prestação de Serviços Advocatícios:"
        )

    # ── Clausula I ────────────────────────────────────────────────

    def _c1_comunicacao(self, doc: Document, data: ContratoConsumidorRequest) -> None:
        doc.add_heading(
            "CLÁUSULA I – DA COMUNICAÇÃO ENTRE AS PARTES – ALERTA DE GOLPES DE FALSOS ADVOGADOS",
            level=2,
        )
        sp = lambda s, p: self._sp(data, s, p)  # noqa: E731

        doc.add_paragraph(
            "1 - " + sp(
                "A CONTRATANTE está ciente e de acordo",
                "Os CONTRATANTES estão cientes e de acordo",
            )
            + " que toda e qualquer comunicação a ser realizada com a CONTRATADA somente "
            "é realizada mediante os canais de atendimento oficiais e abaixo transcritos:"
        )
        # Linhas curtas de contato: justificar esticaria "Email:" ate a margem.
        self._alinhado_a_esquerda(
            doc.add_paragraph(
                "Email: monica@carvalhofurtadoadv.com.br; "
                "gabriela.azevedo@carvalhofurtadoadv.com.br"
            )
        )
        self._alinhado_a_esquerda(doc.add_paragraph("Telefone fixo: (31) 3311-2783"))
        self._alinhado_a_esquerda(
            doc.add_paragraph(
                "Telefone celular: (31) 9.9991-9661 - Monica e (31) 9.8496-7833 – Gabriela"
            )
        )
        doc.add_paragraph(
            "A CONTRATADA não realiza contato através de terceiros e muito menos por meio "
            "de números de telefones desconhecidos."
        )
        doc.add_paragraph(
            "A CONTRATADA não solicita senhas bancárias, não exige a abertura de contas em "
            "bancos, não solicita que "
            + sp(
                "a CONTRATANTE realize procedimentos",
                "os CONTRATANTES realizem procedimentos",
            )
            + " perante suas instituições bancárias, dentre outros dessa natureza."
        )
        doc.add_paragraph(
            "Pagamentos de honorários e despesas estão regidos pelas cláusulas IV e V do "
            "presente contrato de prestação de serviços."
        )
        doc.add_paragraph(
            "A CONTRATADA solicita atenção "
            + sp("da CONTRATANTE", "dos CONTRATANTES")
            + " quanto aos notórios golpes em que dados públicos ou, ainda, dados pessoas "
            "disponibilizados pelo poder público (inclusive tribunais), são utilizados por "
            "golpistas para fazer acreditar na existência de condenações e pagamentos "
            "iminentes nos processos."
        )
        doc.add_paragraph(
            "2 - Caso "
            + sp(
                "a CONTRATANTE receba mensagens ou ligações nesse sentido, desde já se "
                "compromete",
                "os CONTRATANTES recebam mensagens ou ligações nesse sentido, desde já se "
                "comprometem",
            )
            + " a notificar a CONTRATADA."
        )
        doc.add_paragraph(
            "3 - "
            + sp(
                "A CONTRATANTE reconhece e concorda",
                "Os CONTRATANTES reconhecem e concordam",
            )
            + " que a CONTRATADA não se responsabiliza por quaisquer prejuízos, danos ou "
            "perdas decorrentes de contatos, negociações ou contratos realizados por meio de "
            "canais não oficiais ou não autorizados, incluindo, mas não se limitando a "
            "contatos feitos por mensagens, e-mails, telefonemas ou plataformas não oficiais, "
            "reconhecem também que a responsabilidade por eventuais golpes ou fraudes "
            "decorrentes de contatos por meios não autorizados é exclusiva "
            + sp("da CONTRATANTE.", "dos CONTRATANTES.")
        )
        doc.add_paragraph(
            "4 - "
            + sp("A CONTRATANTE está ciente", "Os CONTRATANTES estão cientes")
            + " que qualquer alteração nos meios de comunicação da CONTRATADA será "
            "previamente informado utilizando-se dos meios estabelecidos nesse contrato."
        )
        doc.add_paragraph(
            "5 - Contatos realizados por outros advogados, vinculados à CONTRATADA, serão "
            "sempre ratificados por e-mail com cópia para o e-mail informado nesta cláusula, "
            "devendo "
            + sp("a CONTRATANTE conferir", "os CONTRATANTES conferirem")
            + " os e-mails destinatários e remetentes, inclusive \"Cc\", inclusive se "
            "certificando da presença do e-mail exato indicado nesta cláusula, ou outro que o "
            "venha a substituir, dentre os destinatários."
        )

    # ── Clausula II ───────────────────────────────────────────────

    def _c2_objeto(self, doc: Document, data: ContratoConsumidorRequest) -> None:
        doc.add_heading("CLÁUSULA II – OBJETO", level=2)
        sp = lambda s, p: self._sp(data, s, p)  # noqa: E731

        doc.add_paragraph(
            "1 - Constitui objeto deste contrato a prestação de serviços advocatícios de "
            "contencioso, pela CONTRATADA "
            + sp("a CONTRATANTE", "aos CONTRATANTES")
            + f", perante o {data.juizado}."
        )

        if data.elabora_reclamacao:
            reclamacao = (
                "elaboração de reclamação para que "
                + sp("a CONTRATANTE apresente", "os CONTRATANTES apresentem")
                + " em plataforma de atendimento do consumidor (ex: reclame aqui; "
                "consumidor.gov.br)"
            )
        else:
            reclamacao = (
                "orientação quanto a reclamação extrajudicial a ser apresentada "
                + sp("pela CONTRATANTE", "pelos CONTRATANTES")
                + " em plataforma de atendimento do consumidor (ex: reclame aqui; "
                "consumidor.gov.br)"
            )

        doc.add_paragraph(
            f"1.1 - Os serviços compreendem: atendimento, {reclamacao}; elaboração de "
            "petição inicial, comparecimento em eventual audiência, acompanhamento "
            "processual, eventuais peticionamentos e recursos necessários até o trânsito em "
            "julgado da decisão e execução."
        )
        # Uma ou varias companhias no polo passivo.
        qualificacoes = [
            f"{re.razao_social}, pessoa jurídica de direito privado, inscrita no "
            f"Cadastro Nacional de Pessoas Jurídicas sob o nº {re.cnpj}"
            for re in data.res
        ]
        if len(qualificacoes) == 1:
            rés = qualificacoes[0]
        else:
            rés = ", ".join(qualificacoes[:-1]) + f" e {qualificacoes[-1]}"

        doc.add_paragraph(
            "1.2 – O presente contrato versa sobre a atuação específica relativa à Ação de "
            f"indenização por danos materiais e morais a ser ajuizada em face de {rés}, "
            f"perante o {data.juizado} até decisão final."
        )

    # ── Clausula III ──────────────────────────────────────────────

    def _c3_preco(self, doc: Document, data: ContratoConsumidorRequest) -> None:
        doc.add_heading("CLÁUSULA III – PREÇO:", level=2)
        doc.add_paragraph(
            "Os serviços delimitados na Cláusula II têm o preço fixo estabelecido de 25% "
            "(vinte e cinco por cento) do êxito final da ação, quando do seu recebimento, "
            "inclusive na hipótese de acordo (seja extrajudicialmente, seja judicialmente)"
        )
        doc.add_paragraph(
            "Em caso de aceite de proposta de acordo em que a parte Ré efetue o pagamento "
            "exclusivamente através de milhas aéreas, considerar-se-á o valor do milheiro "
            "(1.000 milhas), para fins de cálculo dos honorários contratuais, conforme tabela "
            "abaixo:"
        )

        # So' as linhas das companhias do caso — nao a tabela inteira do escritorio.
        table = doc.add_table(rows=len(data.res), cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        for i, re in enumerate(data.res):
            table.cell(i, 0).text = re.companhia.upper()
            table.cell(i, 1).text = formatar_valor(re.valor_milheiro)
            # Larguras medidas no CONTRATO.docx: 12,5 cm para a companhia e 3 cm
            # para o valor. Sem fixar, o Word distribui igual e desalinha as linhas.
            table.cell(i, 0).width = Cm(12.5)
            table.cell(i, 1).width = Cm(3)
        self._apply_table_grid(table)

    # ── Clausula IV ───────────────────────────────────────────────

    def _c4_pagamento(self, doc: Document, data: ContratoConsumidorRequest) -> None:
        doc.add_heading("CLÁUSULA IV – PAGAMENTO:", level=2)
        sp = lambda s, p: self._sp(data, s, p)  # noqa: E731

        doc.add_paragraph(
            "1 - Em atenção à cláusula I do presente contrato, a CONTRATADA esclarece que "
            "somente solicitará dados bancários "
            + sp("da CONTRATANTE", "dos CONTRATANTES")
            + " para fins de indicação no processo para confecção de alvará pela secretaria "
            "do Juízo para transferência dos valores objeto da condenação direto "
            + sp("a CONTRATANTE.", "aos CONTRATANTES.")
        )
        doc.add_paragraph(
            "Reitera-se que apenas serão solicitados os dados gerais de contas bancárias "
            "(jamais senhas): banco, agência, número da conta, tipo da conta e beneficiário "
            "com CPF."
        )
        doc.add_paragraph(
            "2 - O percentual estabelecido na cláusula III do contrato deverá ser pago "
            + sp("pela CONTRATANTE", "pelos CONTRATANTES")
            + f" à CONTRATADA em até {data.prazo_pagamento_dias} dias após o recebimento dos "
            "valores/milhas objeto da condenação/acordo, evento que será informado "
            + sp("pela CONTRATANTE", "pelos CONTRATANTES")
            + " à CONTRATADA."
        )
        doc.add_paragraph(
            "3 - Os valores serão depositados diretamente na conta da CONTRATADA abaixo "
            "indicada:"
        )
        for linha in (
            "Banco Bradesco",
            "Agência: 1658",
            "Conta Corrente: 3267-0",
            f"CPF: {CONTRATADA_CPF}",
            "Pix: 31999919661 (celular)",
        ):
            self._alinhado_a_esquerda(doc.add_paragraph(linha))
        doc.add_paragraph(
            "4 - Os honorários ajustados na cláusula III serão devidos, em sua integralidade, "
            "mesmo na hipótese de acordo entre as partes, independentemente se realizados de "
            "forma extrajudicial ou judicial, (desde a assinatura do contrato) ou revogação "
            "de mandato."
        )
        doc.add_paragraph(
            "5 - "
            + sp("A CONTRATANTE está ciente", "Os CONTRATANTES estão cientes")
            + " que os honorários ajustados não se confundem e não são objeto de compensação "
            "com eventuais honorários sucumbenciais advindos da legislação em vigor devidos "
            "pela Empresa Requerida em favor da CONTRATADA no processo."
        )

    # ── Clausula V ────────────────────────────────────────────────

    def _c5_reembolso(self, doc: Document, data: ContratoConsumidorRequest) -> None:
        doc.add_heading("CLÁUSULA V – REEMBOLSO DE DESPESAS", level=2)
        sp = lambda s, p: self._sp(data, s, p)  # noqa: E731

        doc.add_paragraph(
            "1 - O custo de cada cópia xerox, quando não for possível "
            + sp("a CONTRATANTE providenciar", "os CONTRATANTES providenciarem")
            + ", será reembolsado à CONTRATADA no valor de R$ 0,15 (quinze centavos de real)."
        )
        doc.add_paragraph(
            "2 - Despesas referentes ao deslocamento e eventual rotativo/estacionamento "
            "necessários serão reembolsados "
            + sp("pela CONTRATANTE.", "pelos CONTRATANTES.")
        )
        doc.add_paragraph(
            "3 - O reembolso de despesas está condicionado ao envio pela CONTRATADA "
            + sp("a CONTRATANTE", "aos CONTRATANTES")
            + " de relatório contendo os valores de forma discriminada e a que se referem."
        )

    # ── Clausula VI ───────────────────────────────────────────────

    def _c6_condicoes_gerais(self, doc: Document, data: ContratoConsumidorRequest) -> None:
        doc.add_heading("CLÁUSULA VI - CONDIÇÕES GERAIS", level=2)
        sp = lambda s, p: self._sp(data, s, p)  # noqa: E731

        doc.add_paragraph(
            "1 – "
            + sp("A CONTRATANTE", "Os CONTRATANTES")
            + ", sem prejuízo de outras obrigações constantes do presente contrato, se "
            "obrigam:"
        )
        doc.add_paragraph("1.1 - Fornecer todas as informações necessárias para a análise do caso;")
        doc.add_paragraph("1.2 - Disponibilizar e remeter em tempo hábil todos os documentos solicitados;")
        doc.add_paragraph("1.3 - Fazer em dia o pagamento estipulado na Cláusula III;")
        doc.add_paragraph(
            "1.4 - Fornecer e manter atualizada a informação sobre seus dados (endereço e telefone);"
        )
        doc.add_paragraph("1.5 - Comparecer em eventuais audiências designadas;")
        doc.add_paragraph(
            "1.6 - Informar toda e qualquer tentativa de contato direto pela empresa ou quem a represente;"
        )
        doc.add_paragraph(
            "2 - A CONTRATADA se compromete em prestar o serviço jurídico de forma "
            "personalizada e dedicada, não assegurando qualquer chance de êxito nas demandas;"
        )
        doc.add_paragraph(
            "Os pedidos que serão elaborados na Ação são de ciência "
            + sp(
                "da CONTRATANTE que narrou os fatos",
                "dos CONTRATANTES que narraram os fatos",
            )
            + " à CONTRATADA que informou sobre os riscos da demanda, em especial quanto à "
            "sucumbência, bem como improcedência da demanda e do pleito de justiça gratuita."
        )
        doc.add_paragraph(
            "3 - Eventuais custas e emolumentos judiciais, bem como despesas com cálculos, "
            "perícia, honorários sucumbenciais e outros não previstos correrão por conta "
            + sp("da CONTRATANTE.", "dos CONTRATANTES.")
        )
        doc.add_paragraph(
            "4 - Este Contrato só poderá ser alterado, em qualquer de suas disposições, "
            "mediante a celebração, por escrito, de termo aditivo contratual."
        )
        doc.add_paragraph(
            "5 - O presente contrato é título executivo extrajudicial, podendo ser utilizado "
            "para a execução judicial de quaisquer obrigações nele constantes."
        )

    # ── Clausulas VII e VIII ──────────────────────────────────────

    def _c7_prazo(self, doc: Document) -> None:
        doc.add_heading("CLÁUSULA VII – PRAZO", level=2)
        doc.add_paragraph(
            "O prazo de duração dos serviços é por tempo indeterminado, podendo ser "
            "denunciado pelas partes a qualquer tempo mediante aviso escrito com 30 dias de "
            "antecedência, por escrito, ressalvado o direito a percepção dos valores de "
            "honorários acordado pela CONTRATADA."
        )

    def _c8_foro(self, doc: Document, data: ContratoConsumidorRequest) -> None:
        doc.add_heading("CLÁUSULA VIII – FORO", level=2)
        doc.add_paragraph(
            "O foro competente para dirimir os litígios e controvérsias oriundos deste "
            f"contrato é o da Comarca de {data.comarca}."
        )

    # ── Fechamento e assinaturas ──────────────────────────────────

    def _data_contrato(self, data: ContratoConsumidorRequest) -> datetime:
        if not data.data_contrato:
            return datetime.now()
        try:
            return datetime.combine(date.fromisoformat(data.data_contrato[:10]), datetime.min.time())
        except ValueError:
            return datetime.now()

    def _c_fechamento(
        self,
        doc: Document,
        data: ContratoConsumidorRequest,
        signatario_roles: list[dict] | None = None,
    ) -> None:
        doc.add_paragraph(
            "Para os devidos efeitos legais, assinam as partes o presente instrumento em "
            "duas vias de igual teor e forma na presença das testemunhas abaixo."
        )
        doc.add_paragraph()
        doc.add_paragraph(
            f"{data.comarca}, {self._format_date_pt_br(self._data_contrato(data))}."
        )
        doc.add_paragraph()

        # Tags do DocuSeal: mesmos papeis do contrato de honorarios, para o fluxo
        # de assinatura (routers/docuseal.py) funcionar sem mudanca.
        if signatario_roles:
            contratado_sigs = [s for s in signatario_roles if s.get("role", "").startswith("Contratado")]
            advogado_sigs = [s for s in signatario_roles if s.get("role", "").startswith("Advogado")]
            contratante_sigs = [s for s in signatario_roles if s.get("role", "").startswith("Contratante")]
            testemunha_sigs = [s for s in signatario_roles if s.get("role", "").startswith("Testemunha")]

            entries: list[tuple[str, list[str]]] = []
            for sig in contratante_sigs:
                nome = sig.get("name", "Contratante")
                entries.append((
                    f"{{{{Assinatura {nome};type=signature;role={sig['role']}}}}}",
                    [nome.upper(), "Contratante"],
                ))
            # Contratado neste modelo e' sempre a Monica — o fluxo do DocuSeal
            # reaproveita o papel "Contratado" do contrato de honorarios, mas o
            # rotulo no documento nao pode virar "Carvalho & Furtado Advogados".
            for sig in contratado_sigs:
                entries.append((
                    f"{{{{Assinatura {CONTRATADA_NOME};type=signature;role={sig['role']}}}}}",
                    [CONTRATADA_NOME, f"CPF: {CONTRATADA_CPF} - {CONTRATADA_OAB}", "Contratada"],
                ))
            for sig in advogado_sigs:
                nome = sig.get("name", "Advogado")
                entries.append((
                    f"{{{{Assinatura {nome};type=signature;role={sig['role']}}}}}",
                    [nome.upper(), "Advogado"],
                ))
            self._grid_assinaturas(doc, entries)

            if testemunha_sigs:
                doc.add_paragraph()
                doc.add_paragraph("Testemunhas:")
                self._grid_assinaturas(
                    doc,
                    [
                        (
                            f"{{{{Assinatura {s.get('name', 'Testemunha')};type=signature;role={s['role']}}}}}",
                            [f"{s['role'].upper()}: {s.get('name', 'Testemunha').upper()}"],
                        )
                        for s in testemunha_sigs
                    ],
                )
                return
        else:
            entries = []
            for i, c in enumerate(data.contratantes, 1):
                role = "Contratante" if len(data.contratantes) == 1 else f"Contratante {i}"
                if isinstance(c, ContratanteConsumidorPJ):
                    # Quem assina pela empresa e' o representante legal.
                    rotulo = [
                        c.razao_social.upper(),
                        f"CNPJ n.º {c.cnpj}",
                        f"por {c.representante_nome.upper()}",
                        "Contratante",
                    ]
                else:
                    rotulo = [c.nome.upper(), f"CPF n.º {c.cpf}", "Contratante"]
                entries.append((
                    f"{{{{Assinatura {nome_signatario(c)};type=signature;role={role}}}}}",
                    rotulo,
                ))
            entries.append((
                "{{Assinatura Contratada;type=signature;role=Contratado}}",
                [CONTRATADA_NOME, f"CPF: {CONTRATADA_CPF} - {CONTRATADA_OAB}", "Contratada"],
            ))
            self._grid_assinaturas(doc, entries)

        # Testemunhas fisicas (quando nao ha testemunha digital no DocuSeal)
        doc.add_paragraph()
        doc.add_paragraph("Testemunhas:")
        doc.add_paragraph()
        doc.add_paragraph("01 – _______________________    02 – _______________________")
