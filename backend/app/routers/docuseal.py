from __future__ import annotations

import json
import logging
import os
from pathlib import Path
from typing import Any

from fastapi import APIRouter, Depends, Header, HTTPException
from pydantic import BaseModel
from sqlalchemy.orm import Session

from app.auth import CurrentUser, get_current_user
from app.config import BACKEND_DIR, settings
from app.database import AuditLogDB, ContractDB, ContractVersionDB, get_db, utcnow
from app.services.azure_email import AzureEmailService
from app.services.docuseal import DocuSealService

logger = logging.getLogger(__name__)
router = APIRouter(prefix="/api/docuseal", tags=["DocuSeal"])

DOCUSEAL_WEBHOOK_SECRET = os.getenv("DOCUSEAL_WEBHOOK_SECRET", "")


class DocuSealRequest(BaseModel):
    contract_id: str
    signatarios: list[dict[str, str]]


class DocuSealResponse(BaseModel):
    success: bool
    message: str
    submission_id: str | None = None


_docuseal_service: DocuSealService | None = None
_email_service: AzureEmailService | None = None


def get_docuseal_service() -> DocuSealService:
    global _docuseal_service
    if _docuseal_service is None:
        _docuseal_service = DocuSealService()
    return _docuseal_service


def get_email_service() -> AzureEmailService:
    global _email_service
    if _email_service is None:
        _email_service = AzureEmailService()
    return _email_service


def resolve_backend_path(value: str) -> Path:
    path = Path(value)
    if path.is_absolute():
        return path
    return BACKEND_DIR / path


def _resolve_contract_filepath(contract_id: str, db: Session) -> Path:
    """Resolve the contract file path, trying multiple strategies.

    If the file cannot be found on disk (ephemeral filesystem like Render),
    regenerates from form_data_json stored in the database.
    """
    output_dir = resolve_backend_path(settings.output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    # Strategy 1: Get path from DB and check if file exists at that exact location
    latest_ver = (
        db.query(ContractVersionDB)
        .filter(ContractVersionDB.contract_id == contract_id)
        .order_by(ContractVersionDB.version_number.desc())
        .first()
    )

    if latest_ver and latest_ver.file_path:
        stored = Path(latest_ver.file_path)

        # Try the stored path directly
        if stored.exists():
            logger.info("Found contract file at stored path: %s", stored)
            return stored

        # Strategy 2: Try the filename from stored path in current output_dir
        filename = stored.name
        candidate = output_dir / filename
        if candidate.exists():
            logger.info("Found contract file at reconstructed path: %s", candidate)
            return candidate

    # Strategy 3: Convention-based path
    fallback = output_dir / f"contrato_{contract_id}.docx"
    if fallback.exists():
        logger.info("Found contract file at fallback path: %s", fallback)
        return fallback

    # Strategy 4: Regenerate from form_data_json in DB (handles ephemeral filesystems)
    if latest_ver and latest_ver.form_data_json:
        logger.warning(
            "Contract file not found on disk for %s. Regenerating from stored form data...",
            contract_id,
        )
        try:
            import json as _json
            from app.models.contract import ContratoRequest as _CR
            from app.services.contract_generator import ContractGenerator as _CG

            form_data = _json.loads(latest_ver.form_data_json)
            contrato_data = _CR(**form_data)
            gen = _CG()
            _, new_filepath = gen.generate(contrato_data, contract_id=contract_id)
            regenerated = Path(new_filepath)

            # Update the stored path in DB so next time it's found directly
            latest_ver.file_path = str(regenerated)
            db.commit()

            logger.info("Regenerated contract file at: %s", regenerated)
            return regenerated
        except FileNotFoundError as template_err:
            # Template not found (ephemeral FS) - create a minimal placeholder DOCX
            logger.warning("Template not found, creating minimal DOCX: %s", template_err)
            try:
                from docx import Document as _Doc
                doc = _Doc()
                doc.add_paragraph("Contrato em processamento - documento sera regenerado.")
                minimal_path = output_dir / f"contrato_{contract_id}.docx"
                doc.save(str(minimal_path))
                latest_ver.file_path = str(minimal_path)
                db.commit()
                logger.info("Created minimal placeholder DOCX at: %s", minimal_path)
                return minimal_path
            except Exception as min_err:
                logger.error("Failed to create minimal DOCX: %s", min_err)
        except Exception as regen_err:
            logger.error("Failed to regenerate contract %s: %s", contract_id, regen_err)

    # All strategies exhausted
    logger.error(
        "Contract file not found for %s. Tried: stored=%s, output_dir=%s, regeneration=%s",
        contract_id,
        latest_ver.file_path if latest_ver else "N/A",
        output_dir,
        "failed" if latest_ver and latest_ver.form_data_json else "no form data",
    )
    raise HTTPException(status_code=404, detail="Contract file not found")


def _patch_docx_with_signatures(
    existing_filepath: Path,
    signatarios: list[dict],
    contract_id: str,
    db: Session,
    latest_ver,
) -> Path:
    """Patch an existing DOCX to ensure it has signature fields for all signatarios.

    If the existing file can be read, appends missing signature fields.
    If not, creates a new DOCX with all signature fields.
    """
    from docx import Document

    output_dir = resolve_backend_path(settings.output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / f"contrato_{contract_id}.docx"

    # Try to open existing file and patch it
    doc = None
    if existing_filepath.exists():
        try:
            doc = Document(str(existing_filepath))
        except Exception:
            pass

    if doc is None:
        # Create a new minimal document
        doc = Document()
        doc.add_paragraph("CONTRATO DE PRESTACAO DE SERVICOS ADVOCATICIOS")
        doc.add_paragraph()
        doc.add_paragraph("(Documento regenerado para assinatura digital)")
        doc.add_paragraph()

    # Check if signature fields already exist for all roles
    full_text = "\n".join(p.text for p in doc.paragraphs)
    missing_sigs = []
    for sig in signatarios:
        role = sig.get("role", "")
        if f"|signature|{role}" not in full_text:
            missing_sigs.append(sig)

    if missing_sigs:
        # Append signature fields at the end
        if len(doc.paragraphs) > 3:
            doc.add_paragraph()
            doc.add_paragraph("_" * 70)
            doc.add_paragraph()

        # Add ALL signature fields to ensure consistency
        contratado_sigs = [s for s in signatarios if s.get("role", "").startswith("Contratado")]
        advogado_sigs = [s for s in signatarios if s.get("role", "").startswith("Advogado")]
        contratante_sigs = [s for s in signatarios if s.get("role", "").startswith("Contratante")]

        for sig in contratado_sigs:
            role = sig["role"]
            name = sig.get("name", "Contratado")
            doc.add_paragraph(f"{{{{Assinatura {name}|signature|{role}}}}}")
            doc.add_paragraph(f"CONTRATADO: {name.upper()}")
            doc.add_paragraph()

        for sig in advogado_sigs:
            role = sig["role"]
            name = sig.get("name", "Advogado")
            doc.add_paragraph(f"{{{{Assinatura {name}|signature|{role}}}}}")
            doc.add_paragraph(f"ADVOGADO: {name.upper()}")
            doc.add_paragraph()

        for sig in contratante_sigs:
            role = sig["role"]
            name = sig.get("name", "Contratante")
            doc.add_paragraph(f"{{{{Assinatura {name}|signature|{role}}}}}")
            doc.add_paragraph(f"CONTRATANTE: {name.upper()}")
            doc.add_paragraph()

    doc.save(str(output_path))

    # Update DB
    if latest_ver:
        latest_ver.file_path = str(output_path)
        db.commit()

    logger.info("Patched DOCX with %d signature fields at: %s", len(signatarios), output_path)
    return output_path


@router.post("/send-for-signature", response_model=DocuSealResponse)
async def send_for_signature(
    data: DocuSealRequest,
    user: CurrentUser = Depends(get_current_user),
    db: Session = Depends(get_db),
) -> DocuSealResponse:
    """Send contract for digital signature via DocuSeal."""
    try:
        service = get_docuseal_service()

        # Build the full list of signatarios first (need roles to regenerate DOCX)
        all_signatarios = list(data.signatarios)

        # Always include C&F as "Contratado" role (the firm)
        cf_already_included = any(s.get("role") == "Contratado" for s in all_signatarios)
        if not cf_already_included:
            all_signatarios.append({
                "email": settings.cf_signer_email,
                "name": "Carvalho & Furtado Advogados",
                "role": "Contratado",
            })

        # Always include the logged-in lawyer as "Advogado" role
        if user.email:
            user_already_included = any(s.get("email") == user.email for s in all_signatarios)
            if not user_already_included:
                all_signatarios.append({
                    "email": user.email,
                    "name": user.name or user.email,
                    "role": "Advogado",
                })

        # Deduplicate roles: DocuSeal requires unique role per submitter
        # and each role must match a signature field in the template
        from collections import Counter
        role_counts = Counter(s.get("role", "Contratante") for s in all_signatarios)
        role_indices: dict[str, int] = {}
        for sig in all_signatarios:
            role = sig.get("role", "Contratante")
            if role_counts[role] > 1:
                role_indices[role] = role_indices.get(role, 0) + 1
                sig["role"] = f"{role} {role_indices[role]}"

        # Assign order for sequential signing:
        # Contratante(s) sign first, Advogado second, Contratado (C&F) last
        _ROLE_ORDER = {"Contratante": 1, "Advogado": 2, "Contratado": 3}
        for sig in all_signatarios:
            # Extract base role (without number suffix) for order lookup
            base_role = sig.get("role", "Contratante").rstrip(" 0123456789")
            sig["order"] = _ROLE_ORDER.get(base_role, 1)

        # Regenerate DOCX with signature fields matching the unique signatario roles
        latest_ver = (
            db.query(ContractVersionDB)
            .filter(ContractVersionDB.contract_id == data.contract_id)
            .order_by(ContractVersionDB.version_number.desc())
            .first()
        )

        filepath = _resolve_contract_filepath(data.contract_id, db)

        if latest_ver and latest_ver.form_data_json:
            try:
                import json as _json
                from app.models.contract import ContratoRequest as _CR
                from app.services.contract_generator import ContractGenerator as _CG

                form_data = _json.loads(latest_ver.form_data_json)
                contrato_data = _CR(**form_data)
                gen = _CG()
                _, new_filepath = gen.generate(
                    contrato_data,
                    contract_id=data.contract_id,
                    signatario_roles=all_signatarios,
                )
                filepath = Path(new_filepath)
                latest_ver.file_path = str(filepath)
                db.commit()
                logger.info("Regenerated DOCX with matching signature roles for contract %s", data.contract_id)
            except Exception as regen_err:
                logger.warning("Full regeneration failed: %s. Patching DOCX with signature fields.", regen_err)
                # Fallback: patch the existing DOCX or create a new one with signature fields
                try:
                    filepath = _patch_docx_with_signatures(filepath, all_signatarios, data.contract_id, db, latest_ver)
                except Exception as patch_err:
                    logger.error("Failed to patch DOCX with signatures: %s", patch_err)

        result = await service.create_template_from_docx(
            filepath=str(filepath),
            name=f"Contrato Honorarios {data.contract_id}",
        )

        template_id = result.get("id")
        if not template_id:
            return DocuSealResponse(
                success=False,
                message="Failed to create DocuSeal template",
            )

        sign_result = await service.send_for_signature(
            template_id=template_id,
            signatarios=all_signatarios,
            send_email=True,
        )

        if sign_result.get("success"):
            submission = sign_result.get("submission", {})
            submission_id = submission.get("id")

            # Update DB: status + audit log
            contract = db.query(ContractDB).filter(ContractDB.contract_id == data.contract_id).first()
            if contract:
                contract.status = "enviado"
                contract.updated_at = utcnow()
                if latest_ver:
                    latest_ver.docuseal_submission_id = str(submission_id) if submission_id else None
                db.add(AuditLogDB(
                    contract_id=data.contract_id,
                    action="envio_assinatura",
                    detail=f"Enviado para assinatura via DocuSeal (submission {submission_id})",
                    version_number=contract.current_version,
                    user_email=user.email,
                    created_at=utcnow(),
                ))
                db.commit()

            # Send copy of contract to financeiro
            await _send_contract_to_financeiro(data.contract_id, str(filepath), contract, db, user)

            # Send participacao sheet to financeiro if available
            await _send_participacao_to_financeiro(data.contract_id, contract, latest_ver, db, user, contract_filepath=str(filepath))

            return DocuSealResponse(
                success=True,
                message="Documento enviado para assinatura com sucesso",
                submission_id=str(submission_id) if submission_id else None,
            )
        else:
            return DocuSealResponse(
                success=False,
                message=sign_result.get("message", "Erro ao enviar para assinatura"),
            )

    except Exception as e:
        logger.error("DocuSeal error: %s", e)
        raise HTTPException(status_code=500, detail=str(e))


async def _send_contract_to_financeiro(
    contract_id: str,
    filepath: str,
    contract: ContractDB | None,
    db: Session,
    user: CurrentUser,
) -> None:
    """Send a copy of the contract to the financeiro email."""
    try:
        email_service = get_email_service()
        client_name = contract.client_name if contract else "Cliente"

        result = await email_service.send_email_with_attachment(
            to_email=settings.financeiro_email,
            to_name="Financeiro C&F",
            subject=f"Cópia Contrato de Honorários — {client_name}",
            attachment_path=filepath,
            attachment_name=f"contrato_honorarios_{contract_id}.docx",
        )

        if result.get("success"):
            logger.info("Contract copy sent to financeiro for contract %s", contract_id)
            if contract:
                db.add(AuditLogDB(
                    contract_id=contract_id,
                    action="envio_copia_financeiro",
                    detail=f"Cópia do contrato enviada para {settings.financeiro_email}",
                    version_number=contract.current_version,
                    user_email=user.email,
                    created_at=utcnow(),
                ))
                db.commit()
        else:
            logger.warning("Failed to send contract copy to financeiro: %s", result.get("error"))
    except Exception as e:
        logger.error("Error sending contract copy to financeiro: %s", e)


async def _send_participacao_to_financeiro(
    contract_id: str,
    contract: ContractDB | None,
    latest_ver: ContractVersionDB | None,
    db: Session,
    user: CurrentUser,
    contract_filepath: str | None = None,
) -> None:
    """Send participação sheet to financeiro based on stored form data."""
    try:
        if not latest_ver or not latest_ver.form_data_json:
            return

        form_data = json.loads(latest_ver.form_data_json)
        participacao = form_data.get("participacao", {})

        if not participacao.get("tem_participacao"):
            return

        client_name = contract.client_name if contract else "Cliente"

        # Extract "Objeto do Contrato" from escopos - mirror ALL fields the lawyer filled
        ESCOPO_LABELS = {
            "consultoria_contencioso_geral": "Consultoria e contencioso nas areas de atuacao do C&F",
            "contencioso_representacao": "Contencioso para representacao e atuacao em autos",
            "contencioso_memoriais": "Contencioso para Memoriais e sustentacao oral",
            "contencioso_tutela_urgencia": "Contencioso para tutela de urgencia",
            "consultoria_lgpd": "Consultoria LGPD",
            "consultoria_compliance_trabalhista": "Compliance Trabalhista",
            "consultoria_planejamento_tributario": "Planejamento tributario",
            "consultoria_diagnostico_fiscal": "Diagnostico fiscal",
            "consultoria_planejamento_patrimonial": "Planejamento patrimonial",
            "consultoria_estruturacao_societaria": "Estruturacao societaria",
            "consultoria_contratual": "Analise contratual",
            "consultoria_elaboracao_documentos": "Elaboracao de documentos",
            "consultoria_opiniao_legal": "Opiniao legal / Parecer",
            "outro": "",
        }

        escopos = form_data.get("escopos", [])
        escopo_descriptions = []
        for escopo in escopos:
            parts = []
            tipo = escopo.get("tipo_escopo") or escopo.get("tipo", "")
            label = ESCOPO_LABELS.get(tipo, "")

            # Main label (skip for "outro" since descricao_custom will cover it)
            if label and tipo != "outro":
                parts.append(label)

            # Custom description
            desc = escopo.get("escopo_personalizado") or escopo.get("descricao_custom", "")
            if desc:
                parts.append(desc)

            # Process number
            numero_autos = escopo.get("numero_autos", "")
            if numero_autos:
                parts.append(f"Processo: {numero_autos}")

            # Demands
            demandas = escopo.get("demandas", "")
            if demandas:
                parts.append(f"Demandas: {demandas}")

            # People/assets
            pessoas = escopo.get("pessoas_patrimonios", "")
            if pessoas:
                parts.append(f"Pessoas/Patrimonios: {pessoas}")

            # Restructuring type
            reestruturacao = escopo.get("tipo_reestruturacao", "")
            if reestruturacao:
                parts.append(f"Reestruturacao: {reestruturacao}")

            # Documents
            documentos = escopo.get("documentos", "")
            if documentos:
                parts.append(f"Documentos: {documentos}")

            # Legal opinion topic
            consulta = escopo.get("consulta", "")
            if consulta:
                parts.append(f"Consulta: {consulta}")

            # Memorial activities
            subtipo = escopo.get("subtipo_memoriais", {})
            if subtipo:
                atividades = []
                if subtipo.get("elaboracao_memoriais"):
                    atividades.append("Elaboracao de Memoriais")
                if subtipo.get("despacho_memoriais"):
                    atividades.append("Despacho de Memoriais")
                if subtipo.get("sustentacao_oral_relator"):
                    atividades.append("Sustentacao oral c/ Relator")
                if subtipo.get("sustentacao_oral_todos_julgadores"):
                    atividades.append("Sustentacao oral c/ todos os julgadores")
                if atividades:
                    parts.append(f"Atividades: {', '.join(atividades)}")

            if parts:
                escopo_descriptions.append(" | ".join(parts))

        objeto_contrato = "\n".join(escopo_descriptions) if escopo_descriptions else "Nao especificado"

        rows = []
        # Objeto do contrato is the FIRST field (as requested by financeiro)
        # Replace newlines with <br> for HTML rendering
        rows.append(("Objeto do Contrato", objeto_contrato.replace("\n", "<br>")))
        if participacao.get("percentual_ou_valor"):
            rows.append(("Percentual/Valor", participacao["percentual_ou_valor"]))
        if participacao.get("para_quem"):
            rows.append(("Para quem", participacao["para_quem"]))
        if participacao.get("natureza"):
            rows.append(("Natureza", participacao["natureza"]))
        if participacao.get("responsavel_captacao"):
            rows.append(("Resp. Captação", participacao["responsavel_captacao"]))
        if participacao.get("responsavel_gestao"):
            rows.append(("Resp. Gestão", participacao["responsavel_gestao"]))
        if participacao.get("contato_financeiro_cliente"):
            rows.append(("Contato Financeiro Cliente", participacao["contato_financeiro_cliente"]))

        if not rows:
            return

        table_rows = "".join(
            f'<tr><td style="padding:8px;border:1px solid #D7D1CA;font-weight:600;">{k}</td>'
            f'<td style="padding:8px;border:1px solid #D7D1CA;">{v}</td></tr>'
            for k, v in rows
        )

        html = (
            '<div style="font-family: Segoe UI, Tahoma, sans-serif; max-width: 600px;">'
            '<div style="background-color: #1A3C34; padding: 20px 28px; border-radius: 8px 8px 0 0;">'
            '<span style="color: #FFFFFF; font-size: 16px; font-weight: 500;">Ficha de Participação — Uso Interno</span>'
            '</div>'
            '<div style="padding: 24px; border: 1px solid #D7D1CA; border-top: none; border-radius: 0 0 8px 8px;">'
            f'<p><strong>Cliente:</strong> {client_name}</p>'
            f'<p><strong>Contrato:</strong> {contract_id}</p>'
            f'<p><strong>Registrado por:</strong> {user.email}</p>'
            f'<p><strong>Momento:</strong> Envio para assinatura digital</p>'
            '<table style="width:100%;border-collapse:collapse;margin-top:16px;">'
            f'{table_rows}'
            '</table>'
            '</div></div>'
        )

        email_service = get_email_service()

        # Send with contract attachment if filepath is available
        if contract_filepath and Path(contract_filepath).exists():
            result = await email_service.send_html_email_with_attachment(
                to_email=settings.financeiro_email,
                to_name="Financeiro C&F",
                subject=f"Ficha de Participação (Assinatura) — {client_name}",
                html_content=html,
                attachment_path=contract_filepath,
                attachment_name=f"contrato_honorarios_{contract_id}.docx",
            )
        else:
            result = await email_service.send_html_email(
                to_email=settings.financeiro_email,
                to_name="Financeiro C&F",
                subject=f"Ficha de Participação (Assinatura) — {client_name}",
                html_content=html,
            )

        if result.get("success"):
            logger.info("Participacao sheet sent to financeiro for contract %s", contract_id)
            if contract:
                db.add(AuditLogDB(
                    contract_id=contract_id,
                    action="envio_participacao_assinatura",
                    detail=f"Ficha de participação enviada para {settings.financeiro_email} (assinatura)",
                    version_number=contract.current_version,
                    user_email=user.email,
                    created_at=utcnow(),
                ))
                db.commit()
        else:
            logger.warning("Failed to send participacao to financeiro: %s", result.get("error"))
    except Exception as e:
        logger.error("Error sending participacao to financeiro: %s", e)


# ── Webhook (no auth — called externally by DocuSeal) ───────────


class DocuSealWebhookPayload(BaseModel):
    event_type: str
    data: dict[str, Any] = {}


@router.post("/webhook")
async def docuseal_webhook(
    payload: DocuSealWebhookPayload,
    x_docuseal_secret: str | None = Header(None),
    db: Session = Depends(get_db),
):
    """Receive webhook events from DocuSeal (submission.completed / submission.declined)."""
    # Validate webhook secret
    if not DOCUSEAL_WEBHOOK_SECRET or x_docuseal_secret != DOCUSEAL_WEBHOOK_SECRET:
        raise HTTPException(status_code=403, detail="Invalid webhook secret")

    event_type = payload.event_type
    data = payload.data
    submission_id = str(data.get("id", data.get("submission_id", "")))

    logger.info("DocuSeal webhook received: event=%s submission_id=%s", event_type, submission_id)

    if event_type not in ("submission.completed", "submission.declined"):
        return {"status": "ignored", "event_type": event_type}

    new_status = "assinado" if event_type == "submission.completed" else "recusado"

    # Find the contract version that matches this submission
    version = (
        db.query(ContractVersionDB)
        .filter(ContractVersionDB.docuseal_submission_id == submission_id)
        .first()
    )
    if not version:
        logger.warning("No contract version found for submission_id=%s", submission_id)
        return {"status": "not_found", "submission_id": submission_id}

    contract = db.query(ContractDB).filter(ContractDB.contract_id == version.contract_id).first()
    if contract:
        contract.status = new_status
        contract.updated_at = utcnow()
        db.add(AuditLogDB(
            contract_id=contract.contract_id,
            action=f"webhook_{new_status}",
            detail=f"DocuSeal webhook: {event_type} (submission {submission_id})",
            version_number=version.version_number,
            created_at=utcnow(),
        ))
        db.commit()
        logger.info("Contract %s updated to status '%s'", contract.contract_id, new_status)

    return {"status": "ok", "new_status": new_status}


# ── Status check (authenticated) ────────────────────────────────


class DocuSealStatusResponse(BaseModel):
    contract_id: str
    submission_id: str
    status: dict[str, Any]


@router.get("/{contract_id}/status", response_model=DocuSealStatusResponse)
async def get_docuseal_status(
    contract_id: str,
    user: CurrentUser = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Check the signing status of a contract via DocuSeal API."""
    latest_version = (
        db.query(ContractVersionDB)
        .filter(ContractVersionDB.contract_id == contract_id)
        .order_by(ContractVersionDB.version_number.desc())
        .first()
    )

    if not latest_version or not latest_version.docuseal_submission_id:
        raise HTTPException(status_code=404, detail="No DocuSeal submission found for this contract")

    submission_id = latest_version.docuseal_submission_id
    service = get_docuseal_service()
    status_data = await service.get_submission_status(int(submission_id))

    return DocuSealStatusResponse(
        contract_id=contract_id,
        submission_id=submission_id,
        status=status_data,
    )
