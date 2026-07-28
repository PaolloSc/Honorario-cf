from __future__ import annotations
 
import os
import tempfile
import uuid
import zipfile
from copy import deepcopy
from datetime import datetime
from pathlib import Path
 
from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
 
from app.config import BACKEND_DIR, settings
from app.models.contract import (
    Acessorios,
    ContratantePF,
    ContratantePJ,
    ContratoRequest,
    ESCOPO_LABELS,
    EscopoItem,
    Participacao,
    RepresentantePJ,
    SubtipoExito,
    SubtipoMensalidade,
    TipoEscopo,
    TipoHonorario,
    TipoPessoa,
    VariacaoPrecoMensalidade,
)
from app.utils.currency import formatar_valor, valor_com_extenso, valor_por_extenso
 

INCIDENCIA_EXITO_LABELS = {
    "beneficio_economico": "benefício econômico",
    "beneficio_financeiro": "benefício financeiro",
    "beneficio_tributario": "benefício tributário",
    "todos": "benefício econômico, financeiro e/ou tributário",
}

FORMA_PAGAMENTO_LABELS = {
    "a_vista": "à vista",
    "parcelado": "parcelado",
    "conforme_cumprimento": "conforme cumprimento",
}

HONORARIO_LABELS = {
    "pro_labore": "pró-labore",
    "mensalidade": "mensalidade",
    "hora_trabalhada": "hora trabalhada",
}

# numId/abstractNumId da lista multinivel das clausulas (ver _ensure_clause_numbering)
CLAUSE_NUM_ID = 10


class _Numerador:
    """Adiciona as clausulas de um bloco de honorario no nivel certo da lista.

    Bloco unico (sub=False): tudo em x.1, x.2, x.3...
    Varios blocos (sub=True): a 1a clausula e' o chapeu do bloco (x.n) e as
    demais descem um nivel (x.n.1, x.n.2...).
    """

    def __init__(self, gen: "ContractGenerator", doc: Document, sub: bool) -> None:
        self.gen = gen
        self.doc = doc
        self.sub = sub
        self.i = 0

    def __call__(self, texto: str):
        self.i += 1
        ilvl = 2 if (self.sub and self.i > 1) else 1
        return self.gen._add_clausula(self.doc, texto, ilvl)


MESES_PT_BR = [
    "janeiro",
    "fevereiro",
    "março",
    "abril",
    "maio",
    "junho",
    "julho",
    "agosto",
    "setembro",
    "outubro",
    "novembro",
    "dezembro",
]

 
class ContractGenerator:
    def __init__(self) -> None:
        self.template_path = self._resolve_backend_path(settings.template_path)
        self.output_dir = self._resolve_backend_path(settings.output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
 
    def generate(self, data: ContratoRequest, contract_id: str | None = None, signatario_roles: list[dict] | None = None) -> tuple[str, str]:
        """Generate a contract document from form data.

        Returns (contract_id, file_path).
        """
        if contract_id is None:
            contract_id = str(uuid.uuid4())
        doc = self._build_document(data, signatario_roles=signatario_roles)
 
        filename = f"contrato_{contract_id}.docx"
        filepath = self.output_dir / filename
        doc.save(str(filepath))
 
        return contract_id, str(filepath)
 
    def _build_document(self, data: ContratoRequest, signatario_roles: list[dict] | None = None) -> Document:
        doc = self._new_document_from_template()
        self._clear_body(doc)
        self._clear_headers_footers(doc)
        self._ensure_contract_styles(doc)
        self._ensure_clause_numbering(doc)

        self._add_title(doc)
        self._add_parties(doc, data)
        self._add_scope_and_fees(doc, data)
        self._add_fee_details(doc, data)
        self._add_common_clauses(doc, data)
        self._add_accessories(
            doc,
            data.acessorios,
            has_exito=any(TipoHonorario.EXITO in e.honorarios for e in data.escopos),
        )
        self._add_obligations(doc)
        self._add_integrity(doc)
        self._add_term_and_termination(doc, data)
        self._add_ip(doc)
        self._add_general(doc, data)
        self._add_signatures(doc, data, signatario_roles=signatario_roles)
        self._apply_document_standard(doc)
 
        return doc

    def _resolve_backend_path(self, value: str) -> Path:
        path = Path(value)
        if path.is_absolute():
            return path
        return BACKEND_DIR / path

    def _new_document_from_template(self) -> DocxDocument:
        if not self.template_path.exists():
            raise FileNotFoundError(f"Template de contrato nao encontrado: {self.template_path}")

        if self.template_path.suffix.lower() == ".dotx":
            return self._document_from_dotx(self.template_path)

        return Document(str(self.template_path))

    def _document_from_dotx(self, dotx_path: Path) -> DocxDocument:
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        tmp_path = Path(tmp.name)
        tmp.close()

        try:
            with zipfile.ZipFile(dotx_path, "r") as src, zipfile.ZipFile(tmp_path, "w") as dst:
                for item in src.infolist():
                    content = src.read(item.filename)
                    if item.filename == "[Content_Types].xml":
                        content = content.replace(
                            b"application/vnd.openxmlformats-officedocument.wordprocessingml.template.main+xml",
                            b"application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml",
                        )
                    dst.writestr(item, content)

            return Document(str(tmp_path))
        finally:
            try:
                tmp_path.unlink()
            except OSError:
                pass

    def _clear_body(self, doc: DocxDocument) -> None:
        body = doc._element.body
        for element in list(body):
            if element.tag != qn("w:sectPr"):
                body.remove(element)

    def _clear_headers_footers(self, doc: DocxDocument) -> None:
        """Remove all header/footer content (timbrado) from the template."""
        for section in doc.sections:
            for header in (section.header, section.first_page_header, section.even_page_header):
                if header and header._element is not None:
                    for el in list(header._element):
                        header._element.remove(el)
                    header.is_linked_to_previous = False
            for footer in (section.footer, section.first_page_footer, section.even_page_footer):
                if footer and footer._element is not None:
                    for el in list(footer._element):
                        footer._element.remove(el)
                    footer.is_linked_to_previous = False

    def _ensure_contract_styles(self, doc: DocxDocument) -> None:
        if "List Bullet" not in doc.styles:
            bullet_style = doc.styles.add_style("List Bullet", WD_STYLE_TYPE.PARAGRAPH)
            bullet_style.base_style = doc.styles["Normal"]
        else:
            bullet_style = doc.styles["List Bullet"]
        bullet_style.paragraph_format.left_indent = Cm(0)
        bullet_style.paragraph_format.first_line_indent = Cm(0)

        for level in (1, 2, 3):
            style_name = f"Heading {level}"
            if style_name in doc.styles:
                style = doc.styles[style_name]
            else:
                style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)

            style.base_style = doc.styles["Normal"]
            style.font.name = "Segoe UI"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "Segoe UI")
            style.font.size = Pt(12)
            style.font.bold = True
            style.font.color.rgb = RGBColor(0, 0, 0)
            style.paragraph_format.line_spacing = 1.15
            style.paragraph_format.space_after = Pt(6)

    def _ensure_clause_numbering(self, doc: DocxDocument) -> None:
        """Cria a lista multinivel usada pelas seções e cláusulas.

        O Word passa a numerar sozinho (1. / 1.1. / 1.1.1.), então apagar uma
        cláusula renumera as seguintes sem edição manual. O modelo do escritório
        não traz numbering.xml, por isso a parte é criada aqui.
        """
        from docx.opc.constants import CONTENT_TYPE as CT
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
        from docx.opc.packuri import PackURI
        from docx.opc.part import Part

        if any(r.reltype == RT.NUMBERING for r in doc.part.rels.values()):
            return

        def lvl(ilvl: int) -> str:
            fmt = ".".join(f"%{i + 1}" for i in range(ilvl + 1))
            return (
                f'<w:lvl w:ilvl="{ilvl}">'
                '<w:start w:val="1"/><w:numFmt w:val="decimal"/>'
                '<w:suff w:val="space"/>'
                f'<w:lvlText w:val="{fmt}."/><w:lvlJc w:val="left"/>'
                '<w:pPr><w:ind w:left="0" w:firstLine="0"/></w:pPr>'
                "</w:lvl>"
            )

        xml = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<w:numbering xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            f'<w:abstractNum w:abstractNumId="{CLAUSE_NUM_ID}">'
            '<w:multiLevelType w:val="multilevel"/>'
            + "".join(lvl(i) for i in range(3))
            + "</w:abstractNum>"
            f'<w:num w:numId="{CLAUSE_NUM_ID}"><w:abstractNumId w:val="{CLAUSE_NUM_ID}"/></w:num>'
            "</w:numbering>"
        )
        part = Part(
            PackURI("/word/numbering.xml"),
            CT.WML_NUMBERING,
            xml.encode("utf-8"),
            doc.part.package,
        )
        doc.part.relate_to(part, RT.NUMBERING)

    def _numerar(self, paragraph, ilvl: int):
        """Prende o parágrafo à lista de cláusulas no nível informado."""
        num_pr = paragraph._p.get_or_add_pPr().get_or_add_numPr()
        num_pr.get_or_add_ilvl().val = ilvl
        num_pr.get_or_add_numId().val = CLAUSE_NUM_ID
        return paragraph

    def _add_secao(self, doc: Document, titulo: str):
        """Título de seção (nível 0): o número sai automático."""
        return self._numerar(doc.add_heading(titulo, level=2), 0)

    def _add_clausula(self, doc: Document, texto: str, ilvl: int = 1):
        """Cláusula numerada automaticamente (1 = x.y, 2 = x.y.z)."""
        return self._numerar(doc.add_paragraph(texto), ilvl)

    def _apply_document_standard(self, doc: DocxDocument) -> None:
        self._apply_page_setup(doc)
        self._apply_base_styles(doc)
        self._ensure_page_number_footer(doc)

        for paragraph in doc.paragraphs:
            self._format_paragraph(paragraph)

        for table in doc.tables:
            for row in table.rows:
                trPr = row._tr.get_or_add_trPr()
                if trPr.find(qn("w:cantSplit")) is None:
                    trPr.append(OxmlElement("w:cantSplit"))
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._format_paragraph(paragraph)

    def _apply_page_setup(self, doc: DocxDocument) -> None:
        for section in doc.sections:
            section.top_margin = Cm(3)
            section.left_margin = Cm(3)
            section.bottom_margin = Cm(2)
            section.right_margin = Cm(2)

    def _apply_base_styles(self, doc: DocxDocument) -> None:
        for style_name in ("Normal", "Body Text", "List Paragraph"):
            if style_name not in doc.styles:
                continue
            style = doc.styles[style_name]
            style.font.name = "Segoe UI"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "Segoe UI")
            style.font.size = Pt(12)
            style.font.color.rgb = RGBColor(0, 0, 0)
            style.paragraph_format.line_spacing = 1.15
            style.paragraph_format.space_after = Pt(6)

    def _format_paragraph(self, paragraph) -> None:
        is_heading = paragraph.style and paragraph.style.name.startswith("Heading")
        is_signature = self._is_signature_paragraph(paragraph.text)
        is_tag = paragraph.text.strip().startswith("{{")

        paragraph.paragraph_format.line_spacing = 1.15
        paragraph.paragraph_format.space_after = Pt(0 if is_signature else 6)
        paragraph.paragraph_format.keep_together = True
        if is_heading:
            paragraph.paragraph_format.keep_with_next = True

        for run in paragraph.runs:
            if is_heading and not is_tag:
                run.text = run.text.upper()
            run.font.name = "Segoe UI"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Segoe UI")
            run.font.size = Pt(12)
            # Tags DocuSeal ({{...}}) ficam brancas: invisiveis no Word, legiveis p/ o DocuSeal.
            run.font.color.rgb = RGBColor(255, 255, 255) if is_tag else RGBColor(0, 0, 0)
            if is_heading:
                run.bold = True

    def _ensure_page_number_footer(self, doc: DocxDocument) -> None:
        for section in doc.sections:
            footer = section.footer
            if "PAGE" in footer._element.xml:
                continue

            paragraph = footer.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = paragraph.add_run()
            self._add_page_number_field(run)

    def _add_page_number_field(self, run) -> None:
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")

        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = "PAGE"

        separate = OxmlElement("w:fldChar")
        separate.set(qn("w:fldCharType"), "separate")

        text = OxmlElement("w:t")
        text.text = "1"

        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")

        run._r.extend([begin, instr, separate, text, end])

    def _is_signature_paragraph(self, text: str) -> bool:
        stripped = text.strip()
        return (
            stripped.startswith("_")
            or stripped.startswith("{{")
            or stripped in {"Nome:", "CPF:", "TESTEMUNHAS:"}
        )

    def _format_vencimento(self, value: str | None, *, recorrente: bool = False) -> str:
        raw = (value or "").strip()
        if not raw:
            return "a definir"

        digits = "".join(ch for ch in raw if ch.isdigit())
        if recorrente and raw.isdigit():
            return f"no dia {int(raw)} de cada mês"

        if len(digits) == 8:
            return f"em {digits[:2]}/{digits[2:4]}/{digits[4:]}"

        return raw

    def _vencimento_combined(
        self,
        data: str | None,
        obs: str | None,
        legacy: str | None,
        *,
        recorrente: bool = False,
    ) -> str:
        if data:
            try:
                from datetime import datetime as _dt
                dt = _dt.fromisoformat(data)
                if recorrente:
                    base = f"todo dia {dt.day:02d}"
                else:
                    base = f"em {dt.day:02d}/{dt.month:02d}/{dt.year}"
            except Exception:
                base = self._format_vencimento(data, recorrente=recorrente)
            obs_clean = (obs or "").strip()
            if obs_clean:
                return f"{base} ({obs_clean})"
            return base
        if obs and obs.strip():
            return obs.strip()
        return self._format_vencimento(legacy, recorrente=recorrente)

    def _format_percentual(self, value: float) -> str:
        if float(value).is_integer():
            return f"{int(value)}%"
        return f"{value:.2f}".rstrip("0").rstrip(".").replace(".", ",") + "%"

    def _label_from_map(self, value: str | None, labels: dict[str, str]) -> str:
        raw = (value or "").strip()
        return labels.get(raw, raw)

    def _format_date_pt_br(self, value: datetime) -> str:
        return f"{value.day:02d} de {MESES_PT_BR[value.month - 1]} de {value.year}"

    def _apply_table_grid(self, table) -> None:
        tbl_pr = table._tbl.tblPr
        borders = tbl_pr.first_child_found_in("w:tblBorders")
        if borders is None:
            borders = OxmlElement("w:tblBorders")
            tbl_pr.append(borders)

        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            tag = f"w:{edge}"
            element = borders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                borders.append(element)
            element.set(qn("w:val"), "single")
            element.set(qn("w:sz"), "4")
            element.set(qn("w:space"), "0")
            element.set(qn("w:color"), "000000")
 
    def _set_table_header(self, table, *titles: str) -> None:
        """Preenche a linha de titulo da tabela com texto centralizado."""
        for cell, title in zip(table.rows[0].cells, titles):
            cell.text = title
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _add_title(self, doc: Document) -> None:
        title = doc.add_heading("CONTRATO DE PRESTAÇÃO DE SERVIÇOS ADVOCATÍCIOS", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
 
    def _add_parties(self, doc: Document, data: ContratoRequest) -> None:
        self._add_secao(doc, "DAS PARTES")

        unico = len(data.contratantes) == 1
        for i, contratante in enumerate(data.contratantes, 1):
            rotulo = "CONTRATANTE" if unico else f"CONTRATANTE {i}"
            if isinstance(contratante, ContratantePJ) or (
                isinstance(contratante, dict) and contratante.get("tipo") == "PJ"
            ):
                c = contratante if isinstance(contratante, ContratantePJ) else ContratantePJ(**contratante)
                text = (
                    f"{rotulo}: {c.razao_social.upper()}, inscrita no CNPJ n. {self._format_cnpj(c.cnpj)}, "
                    f"com sede em {c.endereco}, e-mail: {c.email}"
                )
                # Qualificacao de cada representante precedida por virgula (nao por ponto).
                quals = [self._qualificar_representante(r) for r in c.representantes if r.nome]
                if quals:
                    text += ", representada por " + "; e por ".join(quals)
                text += "."
            else:
                c = contratante if isinstance(contratante, ContratantePF) else ContratantePF(**contratante)
                qualif = ", ".join(filter(None, [
                    c.nome.upper(),
                    c.nacionalidade,
                    c.profissao,
                    c.estado_civil.value if c.estado_civil else None,
                ]))
                text = (
                    f"{rotulo}: {qualif}, "
                    f"CPF {c.cpf}, residente em {c.endereco}, "
                    f"e-mail: {c.email}."
                )
            doc.add_paragraph(text)

        doc.add_paragraph(
            "CONTRATADO: CARVALHO & FURTADO ADVOGADOS, sociedade simples de advocacia, "
                "inscrita no CNPJ n. 25.463.159/0001-73, com sede na Rua Antônio de Albuquerque, "
                "n. 271, 5º andar, Savassi, Belo Horizonte/MG, a seguir denominado C&F."
        )

    def _qualificar_representante(self, r: "RepresentantePJ") -> str:
        """Qualificacao do representante legal, omitindo campos vazios (sem ', ,')."""
        return ", ".join(filter(None, [
            r.nome.upper(),
            r.nacionalidade,
            r.profissao,
            r.estado_civil.value if r.estado_civil else None,
            f"CPF {r.cpf}" if r.cpf else None,
            f"e-mail: {r.email}" if r.email else None,
        ]))

    def _format_cnpj(self, value: str) -> str:
        """00000000000000 -> 00.000.000/0000-00 (mantem o valor se ja vier formatado)."""
        d = "".join(ch for ch in (value or "") if ch.isdigit())
        if len(d) != 14:
            return value or ""
        return f"{d[:2]}.{d[2:5]}.{d[5:8]}/{d[8:12]}-{d[12:]}"
 
 
    def _add_scope_and_fees(self, doc: Document, data: ContratoRequest) -> None:
        self._add_secao(doc, "OBJETO, ESCOPO E HONORÁRIO")
 
        self._add_clausula(
            doc,
            'O objeto do presente Contrato ("Contrato") é a prestação, pelo C&F, '
            "de serviços advocatícios à CONTRATANTE, conforme o seguinte Escopo e Preço:"
        )
 
        # Summary table
        table = doc.add_table(rows=1, cols=2)
        self._apply_table_grid(table)
        self._set_table_header(table, "Escopo", "Preço")
 
        for escopo in data.escopos:
            row = table.add_row().cells
            row[0].text = self._escopo_description(escopo)
            row[1].text = self._preco_resumo(escopo)
 
        doc.add_paragraph()
 
        self._add_clausula(
            doc,
            "Não estão incluídos no escopo: serviços contábeis, perícias, cálculos, "
            "auditorias, análise econômica, financeira ou de qualquer outra natureza que "
            "não seja estritamente jurídica."
        )
 
        has_hora = any(TipoHonorario.HORA_TRABALHADA in e.honorarios for e in data.escopos)
        has_mensalidade_partido = any(
            TipoHonorario.MENSALIDADE in e.honorarios
            and e.mensalidade
            and e.mensalidade.subtipo == SubtipoMensalidade.ADVOCACIA_PARTIDO
            for e in data.escopos
        )
        has_mensalidade_processo = any(
            TipoHonorario.MENSALIDADE in e.honorarios
            and e.mensalidade
            and e.mensalidade.subtipo in (SubtipoMensalidade.POR_PROCESSO, SubtipoMensalidade.POR_PASTA)
            for e in data.escopos
        )

        if has_hora:
            self._add_clausula(
                doc,
                "Caso a CONTRATANTE solicite atendimento a questões expressamente não "
                "indicadas no objeto e escopo deste Contrato, serão aplicados os mesmos "
                "critérios de honorários de hora trabalhada previstos no Contrato."
            )
        else:
            self._add_clausula(
                doc,
                "Não estão incluídos na precificação serviços oferecidos pelo C&F "
                "e não expressamente indicados no objeto e escopo deste Contrato, os quais "
                "poderão ser pactuados posteriormente entre as Partes."
            )

        if has_mensalidade_partido:
            self._add_clausula(
                doc,
                "Também não estão incluídos na precificação os serviços de consultoria que "
                "constituam um projeto específico multidisciplinar ou dotado de certa "
                "complexidade, tais como planejamentos e/ou estruturações."
            )

        if data.incluir_partes_relacionadas and (has_hora or has_mensalidade_processo):
            self._add_clausula(
                doc,
                "Para fins deste Contrato, são Partes Relacionadas: "
                "(i) cônjuge, companheiro(a) ou parente de primeiro ou segundo grau da "
                "CONTRATANTE; (ii) entidade(s) ou pessoa(s) jurídica(s) cujo controle "
                "fático ou jurídico seja da CONTRATANTE."
            )
            self._add_clausula(
                doc,
                "Caso a CONTRATANTE solicite atendimento a Partes "
                "Relacionadas, salvo ajuste expresso em contrário, serão aplicados os "
                "mesmos critérios de honorários previstos no Contrato, constituindo "
                "nova contratação para todos os fins.",
                ilvl=2,
            )

    def _add_fee_details(self, doc: Document, data: ContratoRequest) -> None:
        self._add_secao(doc, "OUTRAS DISPOSIÇÕES SOBRE HONORÁRIOS")

        # (metodo, objeto) de cada honorario efetivamente preenchido
        blocos: list[tuple] = []
        for escopo in data.escopos:
            for tipo_hon in escopo.honorarios:
                if tipo_hon == TipoHonorario.HORA_TRABALHADA and escopo.hora_trabalhada:
                    blocos.append((self._add_hora_trabalhada, escopo.hora_trabalhada))
                elif tipo_hon == TipoHonorario.PRO_LABORE and escopo.pro_labore:
                    blocos.append((self._add_pro_labore, escopo.pro_labore))
                elif tipo_hon == TipoHonorario.MENSALIDADE and escopo.mensalidade:
                    blocos.append((self._add_mensalidade, escopo.mensalidade))
                elif tipo_hon == TipoHonorario.EXITO and escopo.exito:
                    blocos.append((self._add_exito, escopo.exito))
                elif tipo_hon == TipoHonorario.PERMUTA and escopo.permuta:
                    blocos.append((self._add_permuta, escopo.permuta))

        # Honorario unico: numeracao corrida (3.1, 3.2...) e sem subtitulo.
        # Varios: cada bloco vira 3.n com subclausulas 3.n.1, 3.n.2...
        varios = len(blocos) > 1
        for add, obj in blocos:
            add(doc, obj, _Numerador(self, doc, varios), com_subtitulo=varios)
 
    def _add_hora_trabalhada(self, doc: Document, ht: "HoraTrabalhada", num: "_Numerador", com_subtitulo: bool = False) -> None:
        if com_subtitulo:
            doc.add_heading("HORA TRABALHADA", level=3)
        num(
            f"Em relação ao honorário por hora trabalhada, será observado o seguinte:"
        )
 
        num(
            f"O valor da hora trabalhada será de {valor_com_extenso(ht.valor_hora)}.",
        )

        num(
            f"As horas trabalhadas serão apuradas ao final de cada mês e faturadas em "
            "parcela única no mês imediatamente subsequente.",
        )
 
        if ht.tem_hora_urgencia:
            num(
                f"As horas trabalhadas serão acrescidas de percentual de 50% (cinquenta por cento) "
                "quando, por solicitação da CONTRATANTE, os serviços forem prestados em regime "
                "de urgência.",
            )
 
        if ht.tem_hora_fora_expediente:
            num(
                f"Caso a CONTRATANTE demande a prestação dos serviços após as 19:00 horas ou "
                "durante finais de semana ou feriados, as horas trabalhadas serão acrescidas "
                "do percentual de 100% (cem por cento).",
            )
 
        if ht.tem_hora_urgencia and ht.tem_hora_fora_expediente:
            num(
                f"Caso as horas sejam de urgência e prestadas fora do expediente, serão cobradas "
                "com acréscimo de 150%.",
            )
 
        if ht.tem_teto_mensal and ht.valor_teto_mensal:
            num(
                f"A fatura mensal das horas trabalhadas respeitará o teto de "
                f"{valor_com_extenso(ht.valor_teto_mensal)}, de modo que o valor excedente "
                f"será cobrado na(s) fatura(s) subsequente(s) respeitando o referido teto.",
            )
 
        if ht.tem_pacote_horas and ht.quantidade_horas_pacote and ht.valor_pacote:
            num(
                f"Os serviços jurídicos serão remunerados mediante pacote mensal fixo de "
                f"{ht.quantidade_horas_pacote} horas, no valor de "
                f"{valor_com_extenso(ht.valor_pacote)}.",
            )
            num(
                f"As horas não utilizadas em determinado mês serão acumuladas e poderão ser "
                f"aproveitadas no mês imediatamente subsequente.",
            )
            if ht.duracao_meses:
                num(
                    f"O saldo acumulado será zerado a cada {ht.duracao_meses} meses.",
                )
 
    def _add_pro_labore(self, doc: Document, pl: "ProLabore", num: "_Numerador", com_subtitulo: bool = False) -> None:
        if com_subtitulo:
            doc.add_heading("PRO-LABORE", level=3)
        num(
            f"Em relação ao honorário pró-labore, será observada a seguinte forma de pagamento:"
        )
 
        if pl.tem_parcelamento and pl.numero_parcelas and pl.valor_parcela:
            num(
                f"O valor total de {valor_com_extenso(pl.valor_total)} será pago em "
                f"{pl.numero_parcelas} parcelas de {valor_com_extenso(pl.valor_parcela)}, "
                f"com vencimento {self._vencimento_combined(pl.vencimento_parcelas_data, pl.vencimento_parcelas_obs, pl.vencimento_parcelas, recorrente=True)}."
            )
        else:
            num(
                f"O valor de {valor_com_extenso(pl.valor_total)} será pago em parcela única, "
                f"com vencimento {self._vencimento_combined(pl.vencimento_data, pl.vencimento_obs, pl.vencimento)}."
            )
 
    def _add_mensalidade(self, doc: Document, m: "Mensalidade", num: "_Numerador", com_subtitulo: bool = False) -> None:
        if m.subtipo == SubtipoMensalidade.ADVOCACIA_PARTIDO:
            if com_subtitulo:
                doc.add_heading("MENSALIDADE DE ADVOCACIA DE PARTIDO", level=3)
            num(
                f"Em relação ao honorário por mensalidade de advocacia de partido, "
                "será observado o seguinte:"
            )

            num(
                f"O honorário abrange a prestação de serviços advocatícios de consultoria "
                "e contencioso de rotina nas áreas oferecidas pelo C&F.",
            )

            num(
                f"A precificação possui como referência o fluxo atual de demanda da "
                "CONTRATANTE, sendo que o honorário deverá ser renegociado caso esse "
                "fluxo aumente.",
            )

            num(
                f"O valor da mensalidade será de {valor_com_extenso(m.valor)}.",
            )

            num(
                f"O vencimento da fatura mensal será {self._vencimento_combined(m.dia_vencimento_data, m.dia_vencimento_obs, m.dia_vencimento, recorrente=True)}.",
            )
 
        elif m.subtipo in (SubtipoMensalidade.POR_PROCESSO, SubtipoMensalidade.POR_PASTA):
            tipo_label = "processo" if m.subtipo == SubtipoMensalidade.POR_PROCESSO else "pasta"
            if com_subtitulo:
                doc.add_heading(f"MENSALIDADE POR {tipo_label.upper()}", level=3)
 
            var_label = ""
            if m.variacao_preco == VariacaoPrecoMensalidade.LIMITACAO_TEMPORAL:
                var_label = " com limitação temporal"
            elif m.variacao_preco == VariacaoPrecoMensalidade.REDUCAO_VOLUME:
                var_label = " com redução por volume"
            elif m.variacao_preco == VariacaoPrecoMensalidade.VARIACAO_FASE_PROCESSUAL:
                var_label = " com variação por fase processual"
 
            num(
                f"Em relação ao honorário por mensalidade{var_label}, o vencimento "
                f"da fatura será {self._vencimento_combined(m.dia_vencimento_data, m.dia_vencimento_obs, m.dia_vencimento, recorrente=True)} "
                f"e o valor de {valor_com_extenso(m.valor)} será devido por {tipo_label} enquanto este "
                f"estiver ativo."
            )
 
            if m.variacao_preco == VariacaoPrecoMensalidade.LIMITACAO_TEMPORAL and m.limitacao_temporal_anos:
                num(
                    f"O valor será devido até {m.limitacao_temporal_anos} anos de tramitação "
                    f"sob o patrocínio do C&F."
                )
 
            if m.faixas_preco:
                table = doc.add_table(rows=1, cols=2)
                self._apply_table_grid(table)
                self._set_table_header(table, "Faixa", "Valor")
                for faixa in m.faixas_preco:
                    row = table.add_row().cells
                    row[0].text = faixa.get("faixa", "")
                    row[1].text = faixa.get("valor", "")
 
            num(
                f"Entende-se por ativo aquele {tipo_label} que não foi definitivamente "
                f"extinto, baixado e arquivado no sistema do Tribunal ou respectivo órgão.",
            )
 
    def _add_exito(self, doc: Document, ex: "Exito", num: "_Numerador", com_subtitulo: bool = False) -> None:
        if com_subtitulo:
            doc.add_heading("ÊXITO", level=3)
        num(
            f"Em relação ao honorário de êxito, será observado o seguinte:"
        )
 
        if ex.subtipo == SubtipoExito.PERCENTUAL_FIXO and ex.percentual:
            num(
                f"O percentual de êxito será de {self._format_percentual(ex.percentual)} sobre {ex.base_calculo}."
            )
        elif ex.subtipo == SubtipoExito.PERCENTUAL_VARIAVEL and ex.faixas_percentual:
            num(
                f"O percentual será calculado conforme o valor do Benefício:"
            )
            table = doc.add_table(rows=1, cols=2)
            self._apply_table_grid(table)
            self._set_table_header(table, "Faixa de Valor", "Percentual")
            for faixa in ex.faixas_percentual:
                row = table.add_row().cells
                row[0].text = faixa.get("faixa", "")
                row[1].text = faixa.get("percentual", "")
 
        incidencia = (
            self._label_from_map(ex.incidencia, INCIDENCIA_EXITO_LABELS)
            or "benefício econômico e/ou financeiro e/ou fiscal e/ou tributário"
        )
        num(
            f"O percentual incidirá sobre o {incidencia}, corrigido "
            "monetariamente, aproveitável à CONTRATANTE (Benefício), ainda que parcial.",
        )

        if (ex.forma_pagamento or "").strip():
            num(
                f"Forma de pagamento: {self._label_from_map(ex.forma_pagamento, FORMA_PAGAMENTO_LABELS)}.",
            )

        if ex.vencimento or ex.vencimento_data or ex.vencimento_obs:
            num(
                f"Vencimento: {self._vencimento_combined(ex.vencimento_data, ex.vencimento_obs, ex.vencimento)}.",
            )
 
        if ex.tem_beneficio_prospectivo and ex.prospectivo_duracao_meses:
            num(
                f"Nos casos em que os serviços do C&F também proporcionarem Benefício prospectivo "
                f"à CONTRATANTE, incidirão honorários de êxito calculados sobre o período de "
                f"{ex.prospectivo_duracao_meses} meses.",
            )
 
        if ex.deduz_outro_honorario and ex.honorario_deduzido:
            num(
                f"O honorário de êxito será pago abatendo-se o valor pago a título de "
                f"{self._label_from_map(ex.honorario_deduzido, HONORARIO_LABELS)}.",
            )
 
    def _add_permuta(self, doc: Document, perm: "Permuta", num: "_Numerador", com_subtitulo: bool = False) -> None:
        if com_subtitulo:
            doc.add_heading("PERMUTA", level=3)
        num(
            f"O serviço contratado será permutado com o serviço de {perm.objeto_permuta} "
            f"a ser prestado pela CONTRATANTE ao C&F. {perm.descricao}"
        )
        if perm.tem_torna and perm.valor_torna:
            num(
                f"A torna será de {valor_com_extenso(perm.valor_torna)}, "
                f"paga da seguinte forma: {perm.forma_pagamento_torna or 'a definir'}."
            )
 
    def _add_common_clauses(self, doc: Document, data: ContratoRequest) -> None:
        self._add_secao(doc, "CLÁUSULAS GERAIS")

        has_hora = any(TipoHonorario.HORA_TRABALHADA in e.honorarios for e in data.escopos)
        has_mensalidade_processo = any(
            TipoHonorario.MENSALIDADE in e.honorarios
            and e.mensalidade
            and e.mensalidade.subtipo in (SubtipoMensalidade.POR_PROCESSO, SubtipoMensalidade.POR_PASTA)
            for e in data.escopos
        )
        com_parte_relacionada = data.incluir_partes_relacionadas and (has_hora or has_mensalidade_processo)

        clauses = [
            "Todos os valores previstos nesta contratação serão reajustados anualmente "
            "pela variação positiva e acumulada do IPCA, ou outro índice que vier a "
            "substituí-lo, sempre desde a data da assinatura do Contrato.",
            "Todo e qualquer pagamento devido ao C&F será feito por meio de boleto bancário "
            f"ou transferência bancária para a conta de sua titularidade: {settings.bank_account_info}.",
            "A CONTRATANTE se declara ciente das notórias tentativas gerais de fraude e "
            "golpes simulando contatos de advogados e escritórios de advocacia, estando, "
            "contudo, igualmente ciente dos canais oficiais de contato do C&F e obrigando-se "
            "a realizar pagamentos somente em conta de titularidade do C&F ou mediante "
            "apresentação de boleto ou outro título em que este seja o beneficiário.",
            "A CONTRATANTE reconhece que qualquer pagamento realizado em inobservância ao "
            "previsto neste Contrato será considerado inválido e ineficaz.",
            "As obrigações de pagamento previstas neste Contrato serão devidas, independente "
            "de notificação, tão logo se dê o seu vencimento.",
            "O atraso no pagamento implicará na incidência do seguinte: juros de 1% a.m; "
            "multa de 10% (dez por cento) sobre o valor em atraso e atualização monetária "
            "pelo IPCA, sem prejuízo de suspensão do serviço ou rescisão contratual a "
            "critério do C&F.",
            "Em caso de mudanças legislativas/regulatórias relevantes que alterem a carga "
            "tributária, os custos de conformidade, ou a forma de incidência/retenção de "
            "tributos aplicáveis aos serviços, as Partes renegociarão, de boa-fé, os valores "
            "e/ou a estrutura de faturamento para preservação do equilíbrio "
            "econômico-financeiro.",
            "A CONTRATANTE reconhece que o C&F poderá, dentro da legalidade e das normas "
            "aplicáveis, definir a forma de faturamento mais eficiente do ponto de vista "
            "fiscal (inclusive em eventual migração de regime tributário), sem alteração do "
            "escopo ou do valor líquido pactuado.",
        ]

        if com_parte_relacionada:
            clauses.append(
                "Caso qualificada mais de uma pessoa ou entidade no campo CONTRATANTE, "
                "haverá solidariedade entre elas, assim como no caso de prestação de "
                "serviço a Partes Relacionadas. Na hipótese de obrigações devidas ao C&F, "
                "as Partes reconhecem a possibilidade de encontro de contas, deduções e "
                "compensações ainda que multilaterais entre as partes signatárias e/ou "
                "Partes Relacionadas, de modo a adimplir tais obrigações em ordem "
                "preferencial."
            )
        else:
            clauses.append(
                "Caso qualificada mais de uma pessoa ou entidade no campo CONTRATANTE, "
                "haverá solidariedade entre elas. Na hipótese de obrigações devidas ao "
                "C&F, as Partes reconhecem a possibilidade de encontro de contas, deduções "
                "e compensações ainda que multilaterais entre as partes signatárias, de "
                "modo a adimplir tais obrigações em ordem preferencial."
            )

        for clause in clauses:
            self._add_clausula(doc, clause)
 
    def _add_accessories(self, doc: Document, ac: Acessorios, has_exito: bool = False) -> None:
        self._add_secao(doc, "REEMBOLSOS, DESPESAS E OUTRAS VERBAS")
        if ac.tem_reembolso:
            self._add_clausula(
                doc,
                "Valores adiantados pelo C&F serão reembolsados pela "
                "CONTRATANTE, mediante comprovação, no prazo de até 05 dias após a "
                "apresentação do(s) comprovante(s)."
            )
            if ac.reembolso_limitado and ac.descricao_limitacao_reembolso:
                doc.add_paragraph(f"Limitação: {ac.descricao_limitacao_reembolso}")

        clauses = [
            "Custas, despesas, taxas, emolumentos, cópias xerográficas, diligências, "
            "correspondentes, peritos, assistentes técnicos, tradutores, serviços de "
            "entrega e correio, deslocamentos, transporte, alimentação, hospedagem, demais "
            "despesas necessárias à execução do serviço e eventuais multas processuais e/ou "
            "honorários de sucumbência devidos ao advogado da parte contrária são de "
            "responsabilidade da CONTRATANTE.",
            "A CONTRATANTE reconhece que, para a execução do serviço, o C&F poderá utilizar "
            "ferramentas e/ou sistemas de busca de ativos, endereços e outras informações "
            "como CredLocaliza ou equivalentes, cujo custo será reembolsado pela CONTRATANTE "
            "nos exatos valores faturados pela ferramenta ou sistema.",
            "A prestação de serviço presencial fora da sede do C&F implicará em despesas de "
            f"deslocamento, as quais serão cobradas à razão de {valor_com_extenso(ac.valor_km or 1.70)} "
            "por quilômetro rodado.",
            "O custo de cada cópia xerox a ser reembolsado pela CONTRATANTE é de R$ 0,40 "
            "(quarenta centavos de reais).",
        ]
        # Sucumbencia/renuncia so faz sentido quando ha honorario de exito pactuado.
        if has_exito:
            clauses.append(
                "As Partes pactuam ainda que: (i) em caso de êxito, ainda que parcial, os "
                "honorários sucumbenciais fixados pertencem exclusivamente ao C&F; (ii) em caso "
                "de acordo que inclua renúncia a sucumbências, o C&F deverá ser previamente "
                "consultado; e (iii) se a CONTRATANTE concordar com a redução ou renúncia de "
                "sucumbências sem anuência do C&F, o valor correspondente será descontado do "
                "benefício econômico para fins de cálculo do êxito ou devido diretamente ao C&F."
            )
        for clause in clauses:
            self._add_clausula(doc, clause)
 
    def _add_obligations(self, doc: Document) -> None:
        self._add_secao(doc, "OBRIGAÇÕES DAS PARTES")
        self._add_clausula(
            doc,
            "Obrigações da CONTRATANTE: (i) fornecer informações/documentos de forma "
            "completa e em tempo hábil; (ii) manter dados cadastrais atualizados; (iii) "
            "efetuar pagamentos dentro dos respectivos prazos; (iv) autorizar despesas "
            "quando exigido; (v) cooperar com o C&F na estratégia definida."
        )
        self._add_clausula(
            doc,
            "Obrigações do C&F: (i) executar o serviço com diligência, técnica e zelo; "
            "(ii) manter confidencialidade e sigilo profissional; (iii) fornecer "
            "informações/documentos relativas à prestação de serviços, quando solicitado."
        )
        self._add_clausula(
            doc,
            "A prestação de serviço advocatício constitui obrigação de meio, "
            "inexistindo obrigação de êxito e/ou resultado."
        )
 
    def _add_integrity(self, doc: Document) -> None:
        self._add_secao(doc, "COMPLIANCE")
        self._add_clausula(
            doc,
            "As Partes comprometem-se a observar a legislação aplicável, incluindo Lei "
            "Anticorrupção e outras normas similares, bem como a cooperar com diretrizes de "
            "Governança, quando existentes e conhecidas, no que for pertinente à execução "
            "deste Contrato."
        )
        self._add_clausula(
            doc,
            "As Partes comprometem-se a tratar dados pessoais estritamente para as "
            "finalidades deste Contrato, observando medidas razoáveis de segurança e "
            "confidencialidade, sendo autorizado desde já a criação de cadastros internos "
            "para fins de comunicação em geral."
        )
        self._add_clausula(
            doc,
            "A CONTRATANTE declara estar ciente de que o C&F, sob supervisão humana, "
            "utiliza ferramentas de inteligência artificial e outras tecnologias como "
            "apoio à prestação do serviço."
        )
 
    def _add_term_and_termination(self, doc: Document, data: ContratoRequest) -> None:
        self._add_secao(doc, "PRAZO, RESCISÃO E OUTROS EFEITOS")
        self._add_clausula(
            doc,
            "Ressalvada a hipótese de prazo específico pactuado entre as Partes, o "
            "presente Contrato é celebrado por tempo indeterminado, até que seja esgotado o "
            "objeto contratado."
        )
        self._add_clausula(
            doc,
            "Qualquer Parte poderá rescindir este Contrato imotivadamente mediante "
            "notificação por escrito com antecedência mínima de 30 (trinta) dias."
        )
        self._add_clausula(
            doc,
            "Este prazo de antecedência não substitui nem prejudica o disposto nos "
            "art. 112, §1º, do Código de Processo Civil e 5º, §3º, do Estatuto da OAB, de "
            "modo que, no caso de demandas judiciais, arbitrais ou administrativos, o C&F e "
            "seus advogados permanecerão representando a CONTRATANTE durante os dez dias "
            "seguintes à notificação, salvo se forem substituídos antes do término desse "
            "prazo.",
            ilvl=2,
        )
        base_83 = (
            "Em caso de extinção contratual, aplica-se o seguinte: (i) honorários "
            "vencidos serão devidos integralmente; (ii) honorários vincendos pactuados por "
            "hora trabalhada serão devidos em relação aos serviços executados até a efetiva "
            "extinção; (iii) honorários vincendos pactuados por mensalidade serão devidos "
            "observando-se o prazo de antecedência de 30 dias previstos nesta cláusula; "
            "(iv) honorários vincendos pactuados por pró-labore serão devidos, "
            "proporcionalmente, observando-se os serviços executados e ainda não "
            "remunerados"
        )

        has_exito = any(TipoHonorario.EXITO in e.honorarios for e in data.escopos)
        criterio_exito = (data.acessorios.criterio_extincao_exito or "").strip()
        if has_exito and criterio_exito:
            self._add_clausula(
                doc,
                f"{base_83}; (v) honorários de êxito vincendos ao momento da resilição "
                f"continuarão devidos ao C&F observando-se o seguinte critério: {criterio_exito}."
            )
        elif has_exito:
            self._add_clausula(
                doc,
                f"{base_83}; (v) honorários de êxito vincendos ao momento da resilição "
                "continuarão devidos ao C&F observando-se a seguinte proporção não cumulativa:"
            )
            linhas = [
                ("Antes da primeira decisão de mérito", "50% do percentual de êxito pactuado"),
                ("Depois da primeira decisão de mérito e antes da primeira decisão recursal",
                 "70% do percentual de êxito pactuado"),
                ("Depois da primeira decisão recursal e antes do cumprimento ou liquidação "
                 "definitiva da decisão", "85% do percentual de êxito pactuado"),
                ("Durante cumprimento ou liquidação definitiva da decisão e antes do efetivo "
                 "proveito econômico", "95% do percentual de êxito pactuado"),
                ("Depois do efetivo proveito econômico", "100% do percentual de êxito pactuado"),
            ]
            table = doc.add_table(rows=1, cols=2)
            self._apply_table_grid(table)
            self._set_table_header(
                table, "Fase processual em que for resilido o Contrato", "Honorário devido ao C&F"
            )
            for fase, valor in linhas:
                row = table.add_row().cells
                row[0].text = fase
                row[1].text = valor
            self._add_clausula(
                doc,
                "A eventual inocorrência de determinada fase processual não afeta o "
                "recebimento dos honorários de êxito nos termos previstos nesta cláusula, "
                "aplicando-se o percentual correspondente à fase processual ao tempo da "
                "resilição, independentemente da ocorrência das fases anteriores."
            )
        else:
            self._add_clausula(doc, f"{base_83}.")

        self._add_clausula(
            doc,
            "Exceto se expressa e diversamente pactuado, todas as "
            "disposições contratuais possuem validade e eficácia para os serviços já em "
            "curso."
        )
 
    def _add_ip(self, doc: Document) -> None:
        self._add_secao(doc, "PROPRIEDADE INTELECTUAL")
        self._add_clausula(
            doc,
            "A produção intelectual desenvolvida pelo C&F, como, por exemplo, teses, "
            "estratégias, modelos, documentos, minutas e know-how, permanece de sua "
            "titularidade."
        )
        self._add_clausula(
            doc,
            "Sem expressa autorização do C&F, é vedada a disponibilização a terceiros "
            "do conteúdo dessa produção intelectual (ainda que parcial), ressalvadas "
            "obrigações legais ou ordem de autoridade competente."
        )
        self._add_clausula(
            doc,
            "É facultado ao C&F e aos advogados que o integram valerem-se dessa "
            "produção intelectual em livros, publicações e outras atuações profissionais, "
            "sempre com a ressalva de respeito ao sigilo das questões relacionadas a este "
            "Contrato."
        )
        self._add_clausula(
            doc,
            "A CONTRATANTE autoriza o C&F a utilizar seu nome, marca e logotipo, de "
            "forma não exclusiva, para fins institucionais, inclusive em apresentações, "
            "portfólios e materiais correlatos, sem divulgação de informações confidenciais "
            "do serviço."
        )
 
    def _add_general(self, doc: Document, data: ContratoRequest) -> None:
        self._add_secao(doc, "DISPOSIÇÕES GERAIS")
        gerais = [
            "Será considerada entregue a notificação e/ou comunicação encaminhada ao "
            "endereço declinado no preâmbulo deste Contrato, caso eventual alteração de "
            "contato ou endereço – inclusive eletrônico – não tenha sido devidamente "
            "comunicada ao C&F.",
            "Qualquer termo grafado com letra maiúscula neste Contrato deverá ter o "
            "significado nele previsto.",
            "As Partes se obrigam em caráter irrevogável e irretratável também por seus "
            "sucessores a qualquer título.",
            "Os direitos e obrigações decorrentes deste Contrato não poderão ser cedidos, "
            "salvo com expressa autorização das Partes signatárias.",
            "O não exercício, pelas Partes, de quaisquer dos direitos ou prerrogativas "
            "previstas neste Contrato, ou mesmo na legislação aplicável, será tido como ato "
            "de mera liberalidade, não constituindo alteração ou novação das obrigações ora "
            "estabelecidas, cujo cumprimento poderá ser exigido a qualquer tempo, "
            "independentemente de comunicação prévia à Parte.",
            "As Partes se comprometem a consultar uma à outra sempre que o não-exercício "
            "reiterado de eventual direito trouxer dúvida sobre eventual renúncia tácita, "
            "preferindo a manifestação expressa para a compreensão do comportamento alheio "
            "e formação de legítima confiança.",
            "O presente contrato é título executivo extrajudicial, podendo ser utilizado "
            "para a execução judicial de quaisquer obrigações nele constantes.",
            "Nos termos do artigo 10, § 2º da MP 2200-2/2001, § 4º do artigo 784 do Código "
            "de Processo Civil e legislação correlata, as Partes e as testemunhas aqui "
            "envolvidas reconhecem a validade de assinaturas eletrônicas ainda que não "
            "utilizem de certificado digital emitido pelo padrão ICP-Brasil.",
            "O Contrato terá efeito a partir da data indicada como aquela da sua "
            "formalização, independentemente de as assinaturas, eletrônicas ou não, serem "
            "eventualmente realizadas em data distinta.",
            "Eventual Proposta, feita pelo C&F e aceita pela CONTRATANTE, integra este "
            "Contrato, o qual, no entanto, deverá prevalecer em caso de dúvida, divergência "
            "ou conflito.",
        ]
        for clause in gerais:
            self._add_clausula(doc, clause)

        adicionais = [
            c.strip()
            for c in (data.acessorios.clausulas_adicionais or "").splitlines()
            if c.strip()
        ]
        if adicionais:
            self._add_secao(doc, "DISPOSIÇÕES ADICIONAIS")
            for clause in adicionais:
                self._add_clausula(doc, clause)

        self._add_secao(doc, "FORO")
        self._add_clausula(
            doc,
            "Fica eleito o foro da Comarca de Belo Horizonte/MG para dirimir "
            "quaisquer dúvidas ou controvérsias decorrentes deste Contrato, com renúncia de "
            "qualquer outro, por mais privilegiado que seja."
        )
 
    def _add_signatures(self, doc: Document, data: ContratoRequest, signatario_roles: list[dict] | None = None) -> None:
        doc.add_paragraph()
        doc.add_paragraph(
            f"Belo Horizonte, {self._format_date_pt_br(datetime.now())}."
        )
        doc.add_paragraph()

        if signatario_roles:
            # Generate signature fields matching the exact unique roles from signatarios
            # Group by base role for display ordering: Contratado first, then Advogado(s), then Contratante(s)
            contratado_sigs = [s for s in signatario_roles if s.get("role", "").startswith("Contratado")]
            advogado_sigs = [s for s in signatario_roles if s.get("role", "").startswith("Advogado")]
            contratante_sigs = [s for s in signatario_roles if s.get("role", "").startswith("Contratante")]
            testemunha_sigs = [s for s in signatario_roles if s.get("role", "").startswith("Testemunha")]

            # Partes: dispostas lado a lado (2 por linha)
            partes_entries: list[tuple[str, str]] = []
            for sig in contratado_sigs:
                name = sig.get("name", "Contratado")
                partes_entries.append((
                    f"{{{{Assinatura {name};type=signature;role={sig['role']}}}}}",
                    f"CONTRATADO: {name.upper()}",
                ))
            for sig in advogado_sigs:
                name = sig.get("name", "Advogado")
                partes_entries.append((
                    f"{{{{Assinatura {name};type=signature;role={sig['role']}}}}}",
                    f"ADVOGADO: {name.upper()}",
                ))
            for sig in contratante_sigs:
                name = sig.get("name", "Contratante")
                partes_entries.append((
                    f"{{{{Assinatura {name};type=signature;role={sig['role']}}}}}",
                    f"CONTRATANTE: {name.upper()}",
                ))
            self._render_signature_grid(doc, partes_entries)

            # Testemunhas digitais (assinatura via DocuSeal), tambem lado a lado
            if testemunha_sigs:
                doc.add_paragraph()
                doc.add_paragraph("TESTEMUNHAS:")
                test_entries = [
                    (
                        f"{{{{Assinatura {sig.get('name', 'Testemunha')};type=signature;role={sig['role']}}}}}",
                        f"{sig['role'].upper()}: {sig.get('name', 'Testemunha').upper()}",
                    )
                    for sig in testemunha_sigs
                ]
                self._render_signature_grid(doc, test_entries)
        else:
            # Default: single fields per role (for initial generation without specific signatarios)
            entries: list[tuple[str, str]] = [
                ("{{Assinatura Contratado;type=signature;role=Contratado}}", "CONTRATADO: CARVALHO & FURTADO ADVOGADOS"),
            ]
            for i, contratante in enumerate(data.contratantes, 1):
                if isinstance(contratante, ContratantePJ):
                    nome = contratante.razao_social
                elif isinstance(contratante, ContratantePF):
                    nome = contratante.nome
                else:
                    nome = f"Contratante {i}"

                # Each contratante gets a unique role
                unico = len(data.contratantes) == 1
                role = "Contratante" if unico else f"Contratante {i}"
                rotulo = "CONTRATANTE" if unico else f"CONTRATANTE {i}"
                entries.append((
                    f"{{{{Assinatura {nome};type=signature;role={role}}}}}",
                    f"{rotulo}: {nome.upper()}",
                ))
            self._render_signature_grid(doc, entries)

        # Testemunhas - bloco fisico em branco (somente quando NAO ha testemunhas digitais).
        # Quando signatario_roles inclui papeis "Testemunha", os campos digitais ja foram
        # gerados acima e este bloco e' omitido p/ evitar duplicidade.
        has_digital_testemunhas = bool(
            signatario_roles
            and any(s.get("role", "").startswith("Testemunha") for s in signatario_roles)
        )
        if not has_digital_testemunhas:
            doc.add_paragraph()
            doc.add_paragraph("TESTEMUNHAS:")
            doc.add_paragraph()
            doc.add_paragraph("_" * 50)
            doc.add_paragraph("Nome:")
            doc.add_paragraph("CPF:")
            doc.add_paragraph()
            doc.add_paragraph("_" * 50)
            doc.add_paragraph("Nome:")
            doc.add_paragraph("CPF:")
 
    def _render_signature_grid(self, doc: Document, entries: list[tuple[str, str]]) -> None:
        """Dispoe campos de assinatura lado a lado, 2 por linha, em tabela sem borda.

        Cada `entry` = (tag_assinatura, rotulo). Layout:
            1.    2.
            3.    4.
        """
        if not entries:
            return

        n_rows = (len(entries) + 1) // 2
        table = doc.add_table(rows=n_rows, cols=2)
        table.autofit = True

        for idx, (tag, label) in enumerate(entries):
            cell = table.cell(idx // 2, idx % 2)
            # 1o paragrafo da celula recebe o campo de assinatura (tag DocuSeal).
            # A tag e' escrita em branco: o DocuSeal a le normalmente, mas ela nao
            # polui o DOCX que o advogado abre no Word (aparecia como texto solto).
            cell.paragraphs[0].text = tag
            cell.add_paragraph("_" * 32)
            cell.add_paragraph(label)
            cell.add_paragraph()  # espacamento entre linhas

    def _escopo_description(self, escopo: EscopoItem) -> str:
        if escopo.tipo == TipoEscopo.OUTRO and escopo.descricao_custom:
            return escopo.descricao_custom
 
        label = ESCOPO_LABELS.get(escopo.tipo, str(escopo.tipo))
 
        extras = []
        if escopo.descricao_custom:
            extras.append(escopo.descricao_custom)
        if escopo.numero_autos:
            extras.append(f"nos autos {escopo.numero_autos}")
        if escopo.demandas:
            extras.append(f"para ajuizamento: {escopo.demandas}")
        if escopo.pessoas_patrimonios:
            extras.append(f"pessoas/patrimonios: {escopo.pessoas_patrimonios}")
        if escopo.tipo_reestruturacao:
            extras.append(f"tipo: {escopo.tipo_reestruturacao}")
        if escopo.documentos:
            extras.append(f"documentos: {escopo.documentos}")
        if escopo.consulta:
            extras.append(f"consulta: {escopo.consulta}")
 
        if extras:
            label += " - " + "; ".join(extras)
 
        return label
 
    def _preco_resumo(self, escopo: EscopoItem) -> str:
        parts = []
        for tipo in escopo.honorarios:
            if tipo == TipoHonorario.HORA_TRABALHADA and escopo.hora_trabalhada:
                parts.append(f"{valor_com_extenso(escopo.hora_trabalhada.valor_hora)} por hora trabalhada")
            elif tipo == TipoHonorario.PRO_LABORE and escopo.pro_labore:
                parts.append(f"{valor_com_extenso(escopo.pro_labore.valor_total)} pro-labore")
            elif tipo == TipoHonorario.MENSALIDADE and escopo.mensalidade:
                parts.append(f"{valor_com_extenso(escopo.mensalidade.valor)} de mensalidade")
            elif tipo == TipoHonorario.EXITO and escopo.exito and escopo.exito.percentual:
                parts.append(f"{escopo.exito.percentual}% de exito")
            elif tipo == TipoHonorario.PERMUTA and escopo.permuta:
                parts.append(f"Permuta: {escopo.permuta.objeto_permuta}")
        return " + ".join(parts) if parts else "A definir"
