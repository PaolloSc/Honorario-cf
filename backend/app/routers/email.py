from __future__ import annotations

import logging
import os
import re
import tempfile
from html import escape
from pathlib import Path

from fastapi import APIRouter, Depends, HTTPException
from typing import Annotated

from pydantic import BaseModel, Field, StringConstraints, model_validator
from sqlalchemy.orm import Session

from app.auth import CurrentUser, get_current_user
from app.config import BACKEND_DIR, settings
from app.database import AuditLogDB, ContractDB, ContractVersionDB, get_db, utcnow
from app.routers.contract import _SIG_TAG, _contract_filename
from app.services.azure_email import AzureEmailService

logger = logging.getLogger(__name__)
router = APIRouter(prefix="/api/email", tags=["Email"])


class EmailRequest(BaseModel):
    contract_id: str
    destinatario_email: str
    destinatario_nome: str
    assunto: str = "Contrato de Honorários - C&F Advogados"


class EmailResponse(BaseModel):
    success: bool
    message: str


LOLista = Annotated[
    list[Annotated[str, StringConstraints(max_length=256)]],
    Field(max_length=50),
]


class ParticipacaoEmailRequest(BaseModel):
    contract_id: str
    cliente_nome: str
    objeto_contrato: str = ""
    # Valor estruturado
    valor_tipo: str = ""           # "percentual" | "valor" | "outro"
    valor_percentual: str = ""
    valor_monetario: float | None = None
    valor_outro: str = ""
    # Advogados
    para_quem: list[str] = []
    natureza: str = ""
    responsavel_captacao: str = ""
    responsavel_gestao: str = ""
    # Contato financeiro (3 campos)
    contato_financeiro_nome: str = ""
    contato_financeiro_email: str = ""
    contato_financeiro_telefone: str = ""
    # Base da participacao
    base_tipo: str = ""
    base_escopo_index: int | None = None
    base_honorario: str = ""
    base_label: str = ""
    # Cadastro no Legal One (limites alinhados a coluna legalone_opcoes.valor)
    categoria_cliente: str = Field("", max_length=256)
    etiquetas: LOLista = []
    listas_transmissao: LOLista = []
    # Legados
    percentual_ou_valor: str = ""
    contato_financeiro_cliente: str = ""

    @model_validator(mode="before")
    @classmethod
    def _coerce_nulls(cls, data):
        if isinstance(data, dict):
            for field in ("objeto_contrato", "valor_tipo", "valor_percentual",
                          "valor_outro", "natureza", "responsavel_captacao",
                          "responsavel_gestao", "contato_financeiro_nome",
                          "contato_financeiro_email", "contato_financeiro_telefone",
                          "percentual_ou_valor", "contato_financeiro_cliente",
                          "categoria_cliente"):
                if data.get(field) is None:
                    data[field] = ""
            for field in ("para_quem", "etiquetas", "listas_transmissao"):
                v = data.get(field)
                if isinstance(v, str):
                    data[field] = [v] if v.strip() else []
                elif v is None:
                    data[field] = []
        return data


_email_service: AzureEmailService | None = None


def get_email_service() -> AzureEmailService:
    global _email_service
    if _email_service is None:
        _email_service = AzureEmailService()
    return _email_service


def resolve_backend_path(value: str) -> Path:
    path = Path(value)
    if path.is_absolute():
        return path
    return BACKEND_DIR / path


def _resolve_contract_filepath(contract_id: str, db: Session) -> Path:
    """Resolve the contract file path, trying multiple strategies.

    If the file cannot be found on disk (ephemeral filesystem like Render),
    regenerates from form_data_json stored in the database.
    """
    output_dir = resolve_backend_path(settings.output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    # Strategy 1: Get path from DB and check if file exists at that exact location
    latest_ver = (
        db.query(ContractVersionDB)
        .filter(ContractVersionDB.contract_id == contract_id)
        .order_by(ContractVersionDB.version_number.desc())
        .first()
    )

    if latest_ver and latest_ver.file_path:
        stored = Path(latest_ver.file_path)

        # Try the stored path directly
        if stored.exists():
            logger.info("Found contract file at stored path: %s", stored)
            return stored

        # Strategy 2: Try the filename from stored path in current output_dir
        filename = stored.name
        candidate = output_dir / filename
        if candidate.exists():
            logger.info("Found contract file at reconstructed path: %s", candidate)
            return candidate

    # Strategy 3: Convention-based path
    fallback = output_dir / f"contrato_{contract_id}.docx"
    if fallback.exists():
        logger.info("Found contract file at fallback path: %s", fallback)
        return fallback

    # Strategy 4: Regenerate from form_data_json in DB (handles ephemeral filesystems)
    if latest_ver and latest_ver.form_data_json:
        logger.warning(
            "Contract file not found on disk for %s. Regenerating from stored form data...",
            contract_id,
        )
        try:
            import json as _json
            from app.models.contract import ContratoRequest as _CR
            from app.services.contract_generator import ContractGenerator as _CG

            form_data = _json.loads(latest_ver.form_data_json)
            contrato_data = _CR(**form_data)
            gen = _CG()
            _, new_filepath = gen.generate(contrato_data, contract_id=contract_id)
            regenerated = Path(new_filepath)

            # Update the stored path in DB so next time it's found directly
            latest_ver.file_path = str(regenerated)
            db.commit()

            logger.info("Regenerated contract file at: %s", regenerated)
            return regenerated
        except FileNotFoundError as template_err:
            # Template not found (ephemeral FS) - create a minimal placeholder DOCX
            logger.warning("Template not found, creating minimal DOCX: %s", template_err)
            try:
                from docx import Document as _Doc
                doc = _Doc()
                doc.add_paragraph("Contrato em processamento - documento sera regenerado.")
                minimal_path = output_dir / f"contrato_{contract_id}.docx"
                doc.save(str(minimal_path))
                latest_ver.file_path = str(minimal_path)
                db.commit()
                logger.info("Created minimal placeholder DOCX at: %s", minimal_path)
                return minimal_path
            except Exception as min_err:
                logger.error("Failed to create minimal DOCX: %s", min_err)
        except Exception as regen_err:
            logger.error("Failed to regenerate contract %s: %s", contract_id, regen_err)

    # All strategies exhausted
    logger.error(
        "Contract file not found for %s. Tried: stored=%s, output_dir=%s, regeneration=%s",
        contract_id,
        latest_ver.file_path if latest_ver else "N/A",
        output_dir,
        "failed" if latest_ver and latest_ver.form_data_json else "no form data",
    )
    raise HTTPException(status_code=404, detail="Contract file not found")


def _docx_review_copy(src: Path) -> Path:
    """Cópia de conferência do DOCX para o cliente: troca as merge-tags do DocuSeal
    ({{...;type=signature;...}}) por linhas de assinatura. Não altera o arquivo
    original, que segue com as tags para o fluxo do DocuSeal. Em qualquer falha
    (arquivo não é DOCX válido, erro ao salvar) devolve o próprio `src` — o chamador
    só apaga o resultado quando ele difere de `src`."""
    from docx import Document

    try:
        doc = Document(str(src))
    except Exception as exc:  # arquivo nao e um DOCX valido (ex.: placeholder)
        logger.warning("Nao foi possivel abrir DOCX %s para conferencia: %s", src, exc)
        return src

    def _fix(paras) -> None:
        for p in paras:
            if _SIG_TAG.search(p.text):
                novo = _SIG_TAG.sub("_" * 40, p.text)
                for r in p.runs:
                    r.text = ""
                if p.runs:
                    p.runs[0].text = novo
                else:
                    p.add_run(novo)

    _fix(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                _fix(cell.paragraphs)

    # Caminho único por requisição: evita corrida entre chamadas concorrentes
    # para o mesmo contrato (que compartilhariam review_{stem}.docx).
    fd, tmp_name = tempfile.mkstemp(prefix="review_", suffix=src.suffix or ".docx")
    os.close(fd)
    try:
        doc.save(tmp_name)
    except (OSError, ValueError) as exc:
        logger.warning("Falha ao salvar copia de conferencia de %s: %s", src, exc)
        Path(tmp_name).unlink(missing_ok=True)
        return src
    return Path(tmp_name)


def _tipo_do_contrato(contract_id: str, db: Session) -> str:
    """Define o rotulo do anexo: honorarios ou prestacao de servicos."""
    import json as _json

    from app.services.contract_dispatch import TIPO_HONORARIOS, tipo_contrato

    ver = (
        db.query(ContractVersionDB)
        .filter(ContractVersionDB.contract_id == contract_id)
        .order_by(ContractVersionDB.version_number.desc())
        .first()
    )
    if not ver or not ver.form_data_json:
        return TIPO_HONORARIOS
    try:
        return tipo_contrato(_json.loads(ver.form_data_json))
    except (ValueError, TypeError):
        return TIPO_HONORARIOS


@router.post("/send", response_model=EmailResponse)
async def send_contract_email(
    data: EmailRequest,
    user: CurrentUser = Depends(get_current_user),
    db: Session = Depends(get_db),
) -> EmailResponse:
    """Send contract via email using Azure Communication Services."""
    try:
        filepath = _resolve_contract_filepath(data.contract_id, db)
        contract = db.query(ContractDB).filter(ContractDB.contract_id == data.contract_id).first()

        review_path = _docx_review_copy(filepath)
        attachment_name = _contract_filename(
            contract.client_name if contract else None,
            data.contract_id,
            _tipo_do_contrato(data.contract_id, db),
        )

        service = get_email_service()
        try:
            result = await service.send_email_with_attachment(
                to_email=data.destinatario_email,
                to_name=data.destinatario_nome,
                subject=data.assunto,
                attachment_path=str(review_path),
                attachment_name=attachment_name,
            )
        finally:
            if review_path != filepath:
                review_path.unlink(missing_ok=True)

        if result["success"]:
            if contract:
                db.add(AuditLogDB(
                    contract_id=data.contract_id,
                    action="envio_email",
                    detail=f"E-mail enviado para {data.destinatario_email}",
                    version_number=contract.current_version,
                    user_email=user.email,
                    created_at=utcnow(),
                ))
                db.commit()
            return EmailResponse(success=True, message="E-mail enviado com sucesso")
        else:
            return EmailResponse(success=False, message=result.get("error", "Erro ao enviar e-mail"))

    except HTTPException:
        raise
    except Exception as e:
        logger.error("Failed to send email: %s", e)
        raise HTTPException(status_code=500, detail=str(e))


def _build_objeto_contrato_from_db(contract_id: str, db: Session) -> str:
    """Build 'Objeto do Contrato' text from stored form_data_json as fallback."""
    import json as _json

    ver = (
        db.query(ContractVersionDB)
        .filter(ContractVersionDB.contract_id == contract_id)
        .order_by(ContractVersionDB.version_number.desc())
        .first()
    )
    if not ver or not ver.form_data_json:
        return ""

    try:
        form_data = _json.loads(ver.form_data_json)
    except (ValueError, TypeError):
        return ""

    escopos = form_data.get("escopos", [])
    if not escopos:
        return ""

    from app.models.contract import ESCOPO_LABELS, TipoEscopo

    lines: list[str] = []
    for escopo in escopos:
        parts: list[str] = []
        tipo_raw = escopo.get("tipo", "")

        # Get label from ESCOPO_LABELS enum map
        try:
            tipo_enum = TipoEscopo(tipo_raw)
            label = ESCOPO_LABELS.get(tipo_enum, tipo_raw)
        except ValueError:
            label = tipo_raw

        if label and tipo_raw != "outro":
            parts.append(label)
        if escopo.get("descricao_custom"):
            parts.append(escopo["descricao_custom"])
        if escopo.get("numero_autos"):
            parts.append(f"Processo: {escopo['numero_autos']}")
        if escopo.get("demandas"):
            parts.append(f"Demandas: {escopo['demandas']}")
        if escopo.get("pessoas_patrimonios"):
            parts.append(f"Pessoas/Patrimônios: {escopo['pessoas_patrimonios']}")
        if escopo.get("tipo_reestruturacao"):
            parts.append(f"Reestruturação: {escopo['tipo_reestruturacao']}")
        if escopo.get("documentos"):
            parts.append(f"Documentos: {escopo['documentos']}")
        if escopo.get("consulta"):
            parts.append(f"Consulta: {escopo['consulta']}")

        subtipo_mem = escopo.get("subtipo_memoriais")
        if subtipo_mem:
            atividades: list[str] = []
            if subtipo_mem.get("elaboracao_memoriais"):
                atividades.append("Elaboração de memoriais")
            if subtipo_mem.get("despacho_memoriais"):
                atividades.append("Despacho de memoriais")
            if subtipo_mem.get("sustentacao_oral_relator"):
                atividades.append("Sustentação oral c/ Relator")
            if subtipo_mem.get("sustentacao_oral_todos_julgadores"):
                atividades.append("Sustentação oral c/ todos os julgadores")
            if atividades:
                parts.append(f"Atividades: {', '.join(atividades)}")

        line = " | ".join(parts)
        if line:
            lines.append(line)

    return "\n".join(lines)


@router.post("/send-participacao", response_model=EmailResponse)
async def send_participacao_email(
    data: ParticipacaoEmailRequest,
    user: CurrentUser = Depends(get_current_user),
    db: Session = Depends(get_db),
) -> EmailResponse:
    """Send participation internal sheet to financeiro."""
    try:
        # Use objeto_contrato from request; if empty, build from stored form data
        objeto_contrato = data.objeto_contrato.strip()
        if not objeto_contrato:
            objeto_contrato = _build_objeto_contrato_from_db(data.contract_id, db)

        rows = []
        # Objeto do Contrato first (as requested by financeiro)
        # A quebra de linha vira <br> no laco que monta a tabela, depois do escape.
        if objeto_contrato:
            rows.append(("Objeto do Contrato", objeto_contrato))
        # Base da participacao (escopo ou honorario)
        if data.base_tipo and data.base_label:
            base_prefixo = "Escopo" if data.base_tipo == "escopo" else "Honorário"
            rows.append(("Base", f"{base_prefixo} — {data.base_label}"))
        # Valor (estruturado, com fallback legado)
        if data.valor_tipo == "percentual" and data.valor_percentual:
            rows.append(("Percentual", f"{data.valor_percentual}%"))
        elif data.valor_tipo == "valor" and data.valor_monetario is not None:
            rows.append(("Valor", f"R$ {data.valor_monetario:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")))
        elif data.valor_tipo == "outro" and data.valor_outro:
            rows.append(("Critério", data.valor_outro))
        elif data.percentual_ou_valor:
            rows.append(("Percentual/Valor", data.percentual_ou_valor))
        # Para quem (lista)
        if data.para_quem:
            rows.append(("Para quem", ", ".join(data.para_quem)))
        if data.natureza:
            rows.append(("Natureza", data.natureza))
        if data.responsavel_captacao:
            rows.append(("Resp. Captação", data.responsavel_captacao))
        if data.responsavel_gestao:
            rows.append(("Resp. Gestão", data.responsavel_gestao))
        # Contato financeiro (3 campos, com fallback legado)
        if data.contato_financeiro_nome or data.contato_financeiro_email or data.contato_financeiro_telefone:
            if data.contato_financeiro_nome:
                rows.append(("Contato — Nome", data.contato_financeiro_nome))
            if data.contato_financeiro_email:
                rows.append(("Contato — E-mail", data.contato_financeiro_email))
            if data.contato_financeiro_telefone:
                rows.append(("Contato — Telefone", data.contato_financeiro_telefone))
        elif data.contato_financeiro_cliente:
            rows.append(("Contato Financeiro Cliente", data.contato_financeiro_cliente))
        # Cadastro no Legal One (pode vir sem participacao)
        if data.categoria_cliente:
            rows.append(("Categoria do cliente", data.categoria_cliente))
        if data.etiquetas:
            rows.append(("Etiqueta LO", ", ".join(data.etiquetas)))
        if data.listas_transmissao:
            rows.append(("Lista de transmissão", ", ".join(data.listas_transmissao)))

        # ponytail: participacao inferida dos proprios campos em vez de uma flag nova.
        # O wizard so manda estes campos quando o toggle esta ligado, mas a inferencia
        # cobre todos eles: uma ficha com apenas o responsavel preenchido ainda e
        # participacao, e rotula-la "Cadastro Legal One" seria mentira.
        tem_participacao = any((
            data.base_label,
            data.valor_tipo,
            data.para_quem,
            data.natureza,
            data.responsavel_captacao,
            data.responsavel_gestao,
            data.contato_financeiro_nome,
            data.contato_financeiro_email,
            data.contato_financeiro_telefone,
            data.contato_financeiro_cliente,
            data.percentual_ou_valor,
        ))
        titulo = "Ficha de Participação" if tem_participacao else "Cadastro Legal One"

        # Escapa tudo que veio do payload; a quebra de linha vira <br> depois do
        # escape, para o objeto do contrato continuar legivel sem abrir injecao.
        table_rows = "".join(
            f'<tr><td style="padding:8px;border:1px solid #D7D1CA;font-weight:600;">{escape(k)}</td>'
            f'<td style="padding:8px;border:1px solid #D7D1CA;">{escape(str(v)).replace(chr(10), "<br>")}</td></tr>'
            for k, v in rows
        )

        html = (
            '<div style="font-family: Segoe UI, Tahoma, sans-serif; max-width: 600px;">'
            '<div style="background-color: #1A3C34; padding: 20px 28px; border-radius: 8px 8px 0 0;">'
            f'<span style="color: #FFFFFF; font-size: 16px; font-weight: 500;">{titulo} — Uso Interno</span>'
            '</div>'
            '<div style="padding: 24px; border: 1px solid #D7D1CA; border-top: none; border-radius: 0 0 8px 8px;">'
            f'<p><strong>Cliente:</strong> {escape(data.cliente_nome)}</p>'
            f'<p><strong>Contrato:</strong> {escape(data.contract_id)}</p>'
            f'<p><strong>Registrado por:</strong> {escape(user.email)}</p>'
            '<table style="width:100%;border-collapse:collapse;margin-top:16px;">'
            f'{table_rows}'
            '</table>'
            '</div></div>'
        )

        # Look up contract file for attachment
        contract_filepath = None
        latest_ver = (
            db.query(ContractVersionDB)
            .filter(ContractVersionDB.contract_id == data.contract_id)
            .order_by(ContractVersionDB.version_number.desc())
            .first()
        )
        if latest_ver and latest_ver.file_path:
            stored = Path(latest_ver.file_path)
            if stored.exists():
                contract_filepath = str(stored)
            else:
                # Try convention path
                output_dir = BACKEND_DIR / settings.output_dir
                candidate = output_dir / f"contrato_{data.contract_id}.docx"
                if candidate.exists():
                    contract_filepath = str(candidate)

        service = get_email_service()

        if contract_filepath:
            result = await service.send_html_email_with_attachment(
                to_email=settings.financeiro_email,
                to_name="Financeiro C&F",
                subject=f"{titulo} — {data.cliente_nome}",
                html_content=html,
                attachment_path=contract_filepath,
                attachment_name=_contract_filename(data.cliente_nome, data.contract_id),
            )
        else:
            result = await service.send_html_email(
                to_email=settings.financeiro_email,
                to_name="Financeiro C&F",
                subject=f"{titulo} — {data.cliente_nome}",
                html_content=html,
            )

        if result["success"]:
            contract = db.query(ContractDB).filter(ContractDB.contract_id == data.contract_id).first()
            if contract:
                db.add(AuditLogDB(
                    contract_id=data.contract_id,
                    action="envio_ficha_participacao",
                    detail=f"Ficha de participação enviada para {settings.financeiro_email}",
                    version_number=contract.current_version,
                    user_email=user.email,
                    created_at=utcnow(),
                ))
                db.commit()
            return EmailResponse(success=True, message="Ficha enviada para o financeiro")

        return EmailResponse(success=False, message=result.get("error", "Erro ao enviar ficha"))

    except Exception as e:
        logger.error("Failed to send participacao email: %s", e)
        raise HTTPException(status_code=500, detail=str(e))
