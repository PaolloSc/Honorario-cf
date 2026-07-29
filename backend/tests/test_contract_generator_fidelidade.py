"""Garante que o contrato gerado reproduz o texto do modelo oficial (seções 4-11)."""
import re
import zipfile

from app.models.contract import ContratoRequest
from app.services.contract_generator import ContractGenerator


def _paras_for(req: dict) -> list[str]:
    """Gera o contrato e devolve os parágrafos de texto do .docx."""
    data = ContratoRequest(**req)
    gen = ContractGenerator()
    _, path = gen.generate(data, contract_id="FIDELIDADE_TEST")
    xml = zipfile.ZipFile(path).read("word/document.xml").decode("utf-8")
    paras = []
    for p in re.split(r"</w:p>", xml):
        txt = "".join(re.findall(r"<w:t[^>]*>(.*?)</w:t>", p))
        if txt.strip():
            paras.append(txt.replace("&amp;", "&"))
    return paras


def _base_req(*, honorario="hora_trabalhada", extra_escopo=None, partes_rel=True) -> dict:
    escopo = {"tipo": "consultoria_lgpd", "honorarios": [honorario]}
    if honorario == "hora_trabalhada":
        escopo["hora_trabalhada"] = {
            "valor_hora": 300, "tem_teto_mensal": False, "tem_pacote_horas": False,
            "tem_hora_urgencia": False, "tem_hora_fora_expediente": False,
        }
    if extra_escopo:
        escopo.update(extra_escopo)
    return {
        "contratantes": [{
            "tipo": "PF", "nome": "Fulano", "nacionalidade": "brasileiro",
            "cpf": "00000000000", "profissao": "x", "estado_civil": "Solteiro(a)",
            "endereco": "rua x", "email": "a@a.com",
        }],
        "incluir_partes_relacionadas": partes_rel,
        "escopos": [escopo],
        "acessorios": {"tem_reembolso": True, "reembolso_limitado": False,
                       "tem_penalidade_inadimplemento": False},
        "participacao": {"tem_participacao": False},
    }


def _preview_paras(req: dict) -> list[str]:
    """Parágrafos como o usuário os lê: com a numeração que o Word gera sozinho."""
    from pathlib import Path

    from app.routers.contract import _docx_to_html

    _, path = ContractGenerator().generate(ContratoRequest(**req), contract_id="FIDELIDADE_PREV")
    html = _docx_to_html(Path(path))
    return [
        re.sub(r"<[^>]+>", "", m).replace("&amp;", "&").replace("&quot;", '"')
        for m in re.findall(r"<(?:p|h1|h3)>.*?</(?:p|h1|h3)>", html, re.S)
    ]


def _has(paras: list[str], needle: str) -> bool:
    return any(needle in p for p in paras)


def _escopo(hon: str, *, mensalidade_subtipo="por_processo", mensalidade_variacao="sem_variacao",
            tipo="consultoria_lgpd") -> dict:
    """Escopo válido para um dado tipo de honorário, com o objeto de detalhe exigido."""
    e: dict = {"tipo": tipo, "honorarios": [hon]}
    if hon == "hora_trabalhada":
        e["hora_trabalhada"] = {
            "valor_hora": 300, "tem_teto_mensal": False, "tem_pacote_horas": False,
            "tem_hora_urgencia": False, "tem_hora_fora_expediente": False,
        }
    elif hon == "pro_labore":
        e["pro_labore"] = {"valor_total": 10000, "tem_parcelamento": False}
    elif hon == "mensalidade":
        m = {"valor": 2000, "subtipo": mensalidade_subtipo, "dia_vencimento": "10",
             "variacao_preco": mensalidade_variacao}
        if mensalidade_variacao == "reducao_volume":
            m["faixas_preco"] = [{"faixa": "1-10", "valor": "2000"}, {"faixa": "11+", "valor": "1500"}]
        e["mensalidade"] = m
    elif hon == "exito":
        e["exito"] = {
            "subtipo": "percentual_fixo", "percentual": 20, "incidencia": "beneficio_economico",
            "base_calculo": "x", "vencimento": "a_vista", "forma_pagamento": "x",
            "tem_beneficio_prospectivo": False, "deduz_outro_honorario": False,
        }
    elif hon == "permuta":
        e["permuta"] = {"objeto_permuta": "imovel", "descricao": "x", "tem_torna": False}
    return e


def _req_escopos(escopos: list[dict], *, partes_rel=True) -> dict:
    return {
        "contratantes": [{
            "tipo": "PF", "nome": "Fulano", "nacionalidade": "brasileiro",
            "cpf": "00000000000", "profissao": "x", "estado_civil": "Solteiro(a)",
            "endereco": "rua x", "email": "a@a.com",
        }],
        "incluir_partes_relacionadas": partes_rel,
        "escopos": escopos,
        "acessorios": {"tem_reembolso": True, "reembolso_limitado": False,
                       "tem_penalidade_inadimplemento": False},
        "participacao": {"tem_participacao": False},
    }


def _xml_for(req: dict) -> str:
    data = ContratoRequest(**req)
    gen = ContractGenerator()
    _, path = gen.generate(data, contract_id="FIDELIDADE_XML")
    return zipfile.ZipFile(path).read("word/document.xml").decode("utf-8")


# Casos: cada tipo de honorário, variantes de mensalidade, e combinações.
import pytest
import xml.dom.minidom as _minidom

_CASOS = {
    "hora": [_escopo("hora_trabalhada")],
    "pro_labore": [_escopo("pro_labore")],
    "mensalidade_processo": [_escopo("mensalidade", mensalidade_subtipo="por_processo")],
    "mensalidade_pasta": [_escopo("mensalidade", mensalidade_subtipo="por_pasta")],
    "mensalidade_partido": [_escopo("mensalidade", mensalidade_subtipo="advocacia_partido")],
    "mensalidade_faixas": [_escopo("mensalidade", mensalidade_subtipo="advocacia_partido",
                                   mensalidade_variacao="reducao_volume")],
    "exito": [_escopo("exito")],
    "permuta": [_escopo("permuta")],
    "multi_hora_exito": [_escopo("hora_trabalhada"), _escopo("exito", tipo="consultoria_contratual")],
    "multi_todos": [
        _escopo("hora_trabalhada"),
        _escopo("pro_labore", tipo="consultoria_contratual"),
        _escopo("mensalidade", tipo="consultoria_compliance_trabalhista"),
        _escopo("exito", tipo="contencioso_representacao"),
        _escopo("permuta", tipo="consultoria_opiniao_legal"),
    ],
}


@pytest.mark.parametrize("nome", list(_CASOS.keys()))
def test_secoes_4a11_presentes_para_cada_caso(nome):
    """Cada combinação de honorário/escopo gera todas as seções 4-11 e termina certo."""
    paras = _paras_for(_req_escopos(_CASOS[nome]))
    # Seções comuns sempre presentes (independem do honorário)
    assert _has(paras, "equilíbrio econômico-financeiro")     # 4.7
    assert _has(paras, "CredLocaliza")                       # 5.3
    assert _has(paras, "obrigação de meio")                  # 6.3
    assert _has(paras, "inteligência artificial")            # 7.3
    assert _has(paras, "art. 112, §1º, do Código de Processo Civil")  # 8.2.1
    assert _has(paras, "utilizar seu nome, marca e logotipo")  # 9.4
    assert _has(paras, "título executivo extrajudicial")     # 10.7
    assert _has(paras, "por mais privilegiado que seja")     # 11.1
    assert _has(paras, "TESTEMUNHAS:")                       # termina nas assinaturas


@pytest.mark.parametrize("nome", list(_CASOS.keys()))
def test_xml_bem_formado_para_cada_caso(nome):
    """O .docx gerado é XML bem-formado para toda combinação."""
    xml = _xml_for(_req_escopos(_CASOS[nome]))
    _minidom.parseString(xml)  # levanta se malformado


@pytest.mark.parametrize("nome", ["exito", "multi_hora_exito", "multi_todos"])
def test_tabela_exito_presente_quando_ha_exito(nome):
    paras = _paras_for(_req_escopos(_CASOS[nome]))
    assert _has(paras, "50% do percentual de êxito pactuado")


@pytest.mark.parametrize("nome", ["hora", "pro_labore", "mensalidade_processo", "permuta"])
def test_tabela_exito_ausente_quando_nao_ha_exito(nome):
    paras = _paras_for(_req_escopos(_CASOS[nome]))
    assert not _has(paras, "50% do percentual de êxito pactuado")


@pytest.mark.parametrize("partes_rel", [True, False])
def test_solidariedade_condicional_multi_escopo(partes_rel):
    # hora trabalhada presente => com_parte_relacionada segue incluir_partes_relacionadas
    paras = _paras_for(_req_escopos(_CASOS["multi_todos"], partes_rel=partes_rel))
    tem = _has(paras, "assim como no caso de prestação de serviço a Partes Relacionadas")
    assert tem == partes_rel


def test_secao4_reforma_tributaria_e_fraude_completa():
    paras = _paras_for(_base_req())
    assert _has(paras, "canais oficiais de contato do C&F")
    assert _has(paras, "equilíbrio econômico-financeiro")
    assert _has(paras, "forma de faturamento mais eficiente do ponto de vista fiscal")


def test_secao4_solidariedade_com_parte_relacionada():
    paras = _paras_for(_base_req(partes_rel=True))
    assert _has(paras, "assim como no caso de prestação de serviço a Partes Relacionadas")


def test_secao4_solidariedade_sem_parte_relacionada():
    paras = _paras_for(_base_req(partes_rel=False))
    assert _has(paras, "haverá solidariedade entre elas.")
    assert not _has(paras, "assim como no caso de prestação de serviço a Partes Relacionadas")


def test_secao5_reembolsos_completos():
    paras = _paras_for(_base_req())
    assert _has(paras, "CredLocaliza")
    assert _has(paras, "R$ 1,70")
    assert _has(paras, "R$ 0,40")
    assert _has(paras, "multas processuais e/ou honorários de sucumbência")
    # Sucumbencia/renuncia (5.6) so entra quando ha honorario de exito.
    assert not _has(paras, "honorários sucumbenciais fixados pertencem exclusivamente ao C&F")
    com_exito = _paras_for(_req_escopos([_escopo("exito")]))
    assert _has(com_exito, "honorários sucumbenciais fixados pertencem exclusivamente ao C&F")


def test_secao5_sem_reembolso_omite_51():
    req = _base_req()
    req["acessorios"]["tem_reembolso"] = False
    paras = _paras_for(req)
    assert not _has(paras, "no prazo de até 05 dias")
    assert _has(paras, "CredLocaliza")


def test_secao6_obrigacoes_incisos_completos():
    paras = _paras_for(_base_req())
    assert _has(paras, "autorizar despesas quando exigido")
    assert _has(paras, "cooperar com o C&F na estratégia definida")
    assert _has(paras, "obrigação de meio")


def test_secao7_integridade_lgpd_e_ia():
    paras = _paras_for(_base_req())
    assert _has(paras, "tratar dados pessoais")
    assert _has(paras, "cadastros internos")
    assert _has(paras, "inteligência artificial")
    assert _has(paras, "diretrizes de Governança")


def _req_com_exito() -> dict:
    return _base_req(honorario="exito", extra_escopo={"exito": {
        "subtipo": "percentual_fixo", "percentual": 20, "incidencia": "beneficio_economico",
        "base_calculo": "x", "vencimento": "a_vista", "forma_pagamento": "x",
        "tem_beneficio_prospectivo": False, "deduz_outro_honorario": False,
    }})


def test_secao8_rescisao_cpc_e_extincao():
    paras = _paras_for(_base_req())
    assert _has(paras, "art. 112, §1º, do Código de Processo Civil")
    assert _has(paras, "honorários vencidos serão devidos integralmente")


def test_secao8_tabela_exito_presente_com_exito():
    paras = _paras_for(_req_com_exito())
    assert _has(paras, "50% do percentual de êxito pactuado")
    assert _has(paras, "100% do percentual de êxito pactuado")
    assert _has(paras, "Antes da primeira decisão de mérito")
    assert _has(paras, "inocorrência de determinada fase processual")


def test_secao8_tabela_exito_ausente_sem_exito():
    paras = _paras_for(_base_req(honorario="hora_trabalhada"))
    assert not _has(paras, "50% do percentual de êxito pactuado")


def test_secao9_pi_uso_de_nome_marca():
    paras = _paras_for(_base_req())
    assert _has(paras, "vedada a disponibilização a terceiros")
    assert _has(paras, "utilizar seu nome, marca e logotipo")


def test_secao10_disposicoes_gerais_completas():
    paras = _paras_for(_base_req())
    assert _has(paras, "título executivo extrajudicial")
    assert _has(paras, "MP 2200-2")
    assert _has(paras, "deverá prevalecer em caso de dúvida")


def test_secao11_foro_com_renuncia():
    paras = _paras_for(_base_req())
    assert _has(paras, "com renúncia de qualquer outro, por mais privilegiado que seja")


def test_documento_termina_em_assinaturas():
    paras = _paras_for(_base_req())
    assert _has(paras, "TESTEMUNHAS:")


# ── Feedback 06/07: campos configuráveis e cláusulas ajustadas ──────────────


def test_acessorios_valor_km_configuravel():
    req = _base_req()
    req["acessorios"]["valor_km"] = 2.40
    paras = _paras_for(req)
    assert _has(paras, "R$ 2,40 (dois reais e quarenta centavos)")
    assert not _has(paras, "R$ 1,70")


def test_valor_km_padrao_singular_correto():
    paras = _paras_for(_base_req())
    assert _has(paras, "R$ 1,70 (um real e setenta centavos)")


def test_criterio_extincao_exito_substitui_tabela_de_fases():
    req = _req_com_exito()
    req["acessorios"]["criterio_extincao_exito"] = "assinatura do acordo"
    paras = _preview_paras(req)
    assert _has(paras, "observando-se o seguinte critério: assinatura do acordo.")
    assert not _has(paras, "50% do percentual de êxito pactuado")
    assert not _has(paras, "inocorrência de determinada fase processual")
    assert _has(paras, "8.4. Exceto se expressa")


def test_clausulas_adicionais_numeradas_antes_do_foro():
    req = _base_req()
    req["acessorios"]["clausulas_adicionais"] = "Primeira extra.\nSegunda extra."
    paras = _preview_paras(req)
    assert _has(paras, "DISPOSIÇÕES ADICIONAIS")
    assert _has(paras, "11.1. Primeira extra.")
    assert _has(paras, "11.2. Segunda extra.")
    assert _has(paras, "12.1. Fica eleito o foro")


def test_exito_incidencia_fundida_com_beneficio():
    paras = _paras_for(_req_com_exito())
    assert _has(paras, "incidirá sobre o benefício econômico, corrigido")
    # As antigas 3.5 ("Incidência:") e 3.6 não existem mais separadas
    assert not _has(paras, "Incidência: benefício econômico.")


def test_exito_forma_pagamento_texto_livre():
    req = _req_com_exito()
    req["escopos"][0]["exito"]["forma_pagamento"] = "quando da formalização do acordo"
    paras = _paras_for(req)
    assert _has(paras, "Forma de pagamento: quando da formalização do acordo.")


def test_exito_sem_forma_pagamento_omite_clausula():
    req = _req_com_exito()
    req["escopos"][0]["exito"]["forma_pagamento"] = ""
    paras = _paras_for(req)
    assert not _has(paras, "Forma de pagamento:")


def test_sem_exito_omite_item_v_da_extincao():
    paras = _paras_for(_base_req())
    assert not _has(paras, "(v) honorários de êxito")


def test_linhas_de_tabela_nao_quebram_entre_paginas():
    xml = _xml_for(_req_escopos(_CASOS["exito"]))
    assert "cantSplit" in xml
    assert "keepNext" in xml


def test_preview_nao_vaza_tags_de_assinatura():
    """Feedback: as merge-tags do DocuSeal nao devem aparecer no preview HTML."""
    from pathlib import Path

    from app.routers.contract import _docx_to_html

    data = ContratoRequest(**_base_req())
    _, path = ContractGenerator().generate(data, contract_id="FIDELIDADE_PREVIEW")
    html = _docx_to_html(Path(path))
    assert "type=signature" not in html
    assert "{{" not in html
    # a linha de assinatura e o rotulo do papel continuam visiveis
    assert "____" in html
    assert "CONTRATANTE: FULANO" in html


def test_titulos_de_tabela_centralizados():
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    data = ContratoRequest(**_base_req())
    _, path = ContractGenerator().generate(data, contract_id="FIDELIDADE_HEADER")
    doc = Document(path)
    assert doc.tables, "contrato deveria ter ao menos a tabela Escopo/Preço"
    hdr = doc.tables[0].rows[0].cells
    assert [c.text for c in hdr] == ["Escopo", "Preço"]
    for cell in hdr:
        assert cell.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER


def test_qualificacao_sem_virgula_dupla_com_campo_vazio():
    req = _base_req()
    req["contratantes"][0]["profissao"] = ""  # usuario deixou profissao em branco
    paras = _paras_for(req)
    contratante = next(p for p in paras if p.startswith("CONTRATANTE:"))
    assert ", ," not in contratante
    assert "FULANO, brasileiro, Solteiro(a), CPF" in contratante


def test_preview_titulo_de_tabela_como_th_centralizado():
    from pathlib import Path

    from app.routers.contract import _docx_to_html

    data = ContratoRequest(**_base_req())
    _, path = ContractGenerator().generate(data, contract_id="FIDELIDADE_TH")
    html = _docx_to_html(Path(path))
    assert "<th>Escopo</th><th>Preço</th>" in html
    assert "th{text-align:center}" in html


def _clausulas_secao3(paras: list[str]) -> list[str]:
    """Rotulos (3.x / 3.x.y) das clausulas da secao 3, como saem numerados."""
    return [p.split(" ")[0] for p in paras if re.match(r"^3\.\d", p)]


def test_secao3_honorario_unico_numera_corrido_e_sem_subtitulo():
    paras = _preview_paras(_req_escopos([_escopo("hora_trabalhada")]))
    assert _clausulas_secao3(paras) == ["3.1.", "3.2.", "3.3."]
    assert not _has(paras, "HORA TRABALHADA")


def test_secao3_varios_honorarios_usam_subclausulas():
    paras = _preview_paras(_req_escopos([_escopo("hora_trabalhada"), _escopo("pro_labore")]))
    # 1o bloco: 3.1 (chapeu) + 3.1.1, 3.1.2; 2o bloco: 3.2 (chapeu) + 3.2.1
    assert _clausulas_secao3(paras) == ["3.1.", "3.1.1.", "3.1.2.", "3.2.", "3.2.1."]
    assert _has(paras, "HORA TRABALHADA") and _has(paras, "PRO-LABORE")


def test_clausulas_sao_numeradas_pelo_word_e_nao_no_texto():
    """O numero vem da lista multinivel do Word: apagar uma clausula renumera as demais."""
    import zipfile

    from app.services.contract_generator import CLAUSE_NUM_ID

    data = ContratoRequest(**_base_req())
    _, path = ContractGenerator().generate(data, contract_id="FIDELIDADE_NUM")
    z = zipfile.ZipFile(path)
    assert "word/numbering.xml" in z.namelist()

    xml = z.read("word/document.xml").decode("utf-8")
    numeradas = 0
    for p in re.split(r"</w:p>", xml):
        txt = "".join(re.findall(r"<w:t[^>]*>(.*?)</w:t>", p)).strip()
        if f'<w:numId w:val="{CLAUSE_NUM_ID}"' in p:
            numeradas += 1
        # Varre TODO paragrafo, nao so os da lista: uma clausula numerada a mao
        # fica FORA da lista, e um teste que olhasse so os itens da lista nunca
        # a veria — foi assim que esta assercao passou com o bug de volta.
        assert not re.match(r"^\d+\.\d*\.? \S", txt), f"numero escrito a mao: {txt[:60]}"
    assert numeradas > 20


def test_preview_nao_junta_rotulo_com_a_linha_de_assinatura():
    """Na celula, cada paragrafo e' uma linha: o rotulo nao pode colar nos underscores."""
    from pathlib import Path

    from app.routers.contract import _docx_to_html

    data = ContratoRequest(**_base_req())
    _, path = ContractGenerator().generate(
        data,
        contract_id="FIDELIDADE_SIG",
        signatario_roles=[
            {"email": "a@a.com", "name": "Fulano", "role": "Contratante"},
            {"email": "cf@cf.br", "name": "Carvalho & Furtado Advogados", "role": "Contratado"},
        ],
    )
    html = _docx_to_html(Path(path))
    assert "<br>CONTRATANTE: FULANO" in html

    # Cada celula de assinatura tem UMA linha de underscores. A previa apagava a
    # tag do DocuSeal trocando-a por underscores, o que somava uma segunda linha;
    # conferir so o par "underscores<br>rotulo" nao pegava isso, porque a linha
    # duplicada fica ANTES e o par continuava batendo.
    celulas = re.findall(r"<td>(.*?)</td>", html, re.S)
    assinatura = [c for c in celulas if "CONTRATANTE:" in c or "CONTRATADO:" in c]
    assert assinatura, "nenhuma celula de assinatura no preview"
    for celula in assinatura:
        linhas_underscore = [
            linha for linha in celula.split("<br>") if set(linha.strip()) == {"_"}
        ]
        assert len(linhas_underscore) == 1, f"esperava 1 linha de assinatura: {celula[:120]}"


def _assinatura_pPr(path: str) -> list[tuple[str, str, str, str]]:
    """(texto, alinhamento, espaco_antes, espaco_depois) de cada paragrafo da grade."""
    import zipfile

    xml = zipfile.ZipFile(path).read("word/document.xml").decode("utf-8")
    grade = re.findall(r"<w:tbl>.*?</w:tbl>", xml, re.S)[-1]
    out = []
    for p in re.split(r"</w:p>", grade):
        if not "".join(re.findall(r"<w:t[^>]*>(.*?)</w:t>", p)).strip():
            continue
        jc = re.search(r'<w:jc w:val="(\w+)"/>', p)
        antes = re.search(r'w:before="(\d+)"', p)
        depois = re.search(r'w:after="(\d+)"', p)
        out.append((
            "".join(re.findall(r"<w:t[^>]*>(.*?)</w:t>", p)).strip(),
            jc.group(1) if jc else "herdado",
            antes.group(1) if antes else "herdado",
            depois.group(1) if depois else "herdado",
        ))
    return out


def test_assinatura_centralizada_na_grade_e_testemunhas_a_esquerda():
    """Dois alinhamentos distintos, confirmados no .docx que o escritorio ajustou.

    O modelo justifica tudo (w:jc=both) e espalhava "CONTRATANTE: MARINA ALVES
    RIBEIRO" na largura da celula, entao o alinhamento precisa ser fixado. Na
    grade e' centralizado dentro da celula; o bloco de testemunhas, que corre
    solto no corpo, fica a esquerda como o resto do contrato. Centralizar os
    dois foi recusado pelo escritorio.
    """
    from app.services.contract_generator import ESPACO_ASSINATURA

    data = ContratoRequest(**_base_req())
    _, path = ContractGenerator().generate(
        data,
        contract_id="FIDELIDADE_FMT",
        signatario_roles=[
            {"email": "a@a.com", "name": "Fulano", "role": "Contratante"},
            {"email": "cf@cf.br", "name": "Carvalho & Furtado Advogados", "role": "Contratado"},
        ],
    )
    folga = str(int(ESPACO_ASSINATURA.pt * 20))

    grade = _assinatura_pPr(path)
    assert grade, "grade de assinaturas vazia"
    for texto, jc, antes, depois in grade:
        assert jc == "center", f"grade deveria ser centralizada, veio {jc} em {texto[:30]!r}"
        assert depois == "0", f"espaco sobrando depois de {texto[:30]!r}: {depois}"
        # A folga fica so acima da linha de assinatura; o nome cola nela.
        assert antes == (folga if set(texto) == {"_"} else "0"), texto[:30]

    testemunhas = _testemunhas_pPr(path)
    assert testemunhas, "bloco de testemunhas vazio"
    for texto, jc, antes, _depois in testemunhas:
        assert jc == "left", f"testemunhas deveriam ficar a esquerda, veio {jc}"
        assert antes == (folga if set(texto) == {"_"} else "0"), texto[:30]


def _testemunhas_pPr(path: str) -> list[tuple[str, str, str, str]]:
    """Mesma leitura de _assinatura_pPr, para o bloco solto no corpo."""
    import zipfile

    xml = zipfile.ZipFile(path).read("word/document.xml").decode("utf-8")
    corpo = xml.split("</w:tbl>")[-1]
    out = []
    for p in re.split(r"</w:p>", corpo):
        texto = "".join(re.findall(r"<w:t[^>]*>(.*?)</w:t>", p)).strip()
        if not texto:
            continue
        jc = re.search(r'<w:jc w:val="(\w+)"/>', p)
        antes = re.search(r'w:before="(\d+)"', p)
        depois = re.search(r'w:after="(\d+)"', p)
        out.append((
            texto,
            jc.group(1) if jc else "herdado",
            antes.group(1) if antes else "herdado",
            depois.group(1) if depois else "herdado",
        ))
    return out
