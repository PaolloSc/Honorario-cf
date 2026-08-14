"""Checagem do contrato de consumidor: concordancia singular/plural, milheiro e Ré.

Reproduz os dois modelos do escritorio: CONTRATO 2025 (1 contratante, KLM) e
CONTRATO (2 contratantes, Azul).
"""

from docx import Document

from app.models.contrato_consumidor import ContratoConsumidorRequest
from app.services.consumidor_generator import ConsumidorGenerator


def _texto(filepath: str) -> str:
    doc = Document(filepath)
    partes = [p.text for p in doc.paragraphs]
    for tabela in doc.tables:
        for linha in tabela.rows:
            partes.extend(c.text for c in linha.cells)
    return "\n".join(partes)


UM = {
    "contratantes": [
        {
            "nome": "Maria de Fátima Soares Pereira",
            "genero": "F",
            "cpf": "761.794.436-53",
            "rg": "MG-3.764.023",
            "endereco": "Rua Pio Porto de Menezes, n.º 115, apto. 803, Bairro Luxemburgo, Belo Horizonte/MG, CEP: 30.380-300",
            "email": "mariafatimasp@gmail.com",
            "celular": "(31) 9972-1792",
        }
    ],
    "companhia": "KLM",
    "re_razao_social": "KLM CIA REAL HOLANDESA DE AVIACAO",
    "re_cnpj": "33.643.420/0001-45",
    "juizado": "Juizado Especial Cível de Belo Horizonte - MG",
    "elabora_reclamacao": True,
}

DOIS = {
    "contratantes": [
        {
            "nome": "Daniela Elias Araujo Marques",
            "genero": "F",
            "cpf": "033.395.236-73",
            "endereco": "Rua Major Egido Luiz Cerqueira, n.º 155E, Interfone 05, Centro, Itapecerica/MG, CEP: 35.550-000",
        },
        {
            "nome": "Henrique de Oliveira Marques",
            "genero": "M",
            "cpf": "072.029.566-16",
            "endereco": "Rua Major Egido Luiz Cerqueira, n.º 155E, Interfone 05, Centro, Itapecerica/MG, CEP: 35.550-000",
        },
    ],
    "companhia": "Azul",
    "re_razao_social": "AZUL LINHAS AEREAS BRASILEIRAS S.A.",
    "re_cnpj": "09.296.295/0001-60",
    "prazo_pagamento_dias": 5,
    "elabora_reclamacao": False,
}


PJ = {
    "tipo": "PJ",
    "razao_social": "Transportadora Silva Ltda",
    "cnpj": "12.345.678/0001-90",
    "endereco": "Av. Afonso Pena, n.º 1000, Centro, Belo Horizonte/MG, CEP: 30130-000",
    "email": "contato@silva.com.br",
    "representante_nome": "Carlos Silva",
    "representante_cpf": "111.222.333-44",
    "representante_genero": "M",
    "representante_email": "carlos@silva.com.br",
}


def test_pj_sozinha_qualifica_com_representante(tmp_path):
    gen = ConsumidorGenerator()
    gen.output_dir = tmp_path
    texto = _texto(gen.generate(ContratoConsumidorRequest(**{**UM, "contratantes": [PJ]}))[1])

    assert "TRANSPORTADORA SILVA LTDA, pessoa jurídica de direito privado" in texto
    assert "inscrita no CNPJ sob o n.º 12.345.678/0001-90" in texto
    assert "com sede à Av. Afonso Pena" in texto
    assert "neste ato representada por CARLOS SILVA, brasileiro, inscrito no CPF" in texto
    assert "doravante denominada apenas CONTRATANTE" in texto
    # Empresa nao tem domicilio no texto do contrato.
    assert "residente e domiciliad" not in texto
    # Quem assina pela empresa e' o representante.
    assert "Assinatura Carlos Silva" in texto
    assert "por CARLOS SILVA" in texto


def test_pf_e_pj_juntos(tmp_path):
    """Empresa concorda no feminino: com uma contratante mulher, o grupo fica feminino."""
    gen = ConsumidorGenerator()
    gen.output_dir = tmp_path
    dados = {**UM, "contratantes": [UM["contratantes"][0], PJ]}
    texto = _texto(gen.generate(ContratoConsumidorRequest(**dados))[1])

    assert "doravante denominadas apenas CONTRATANTES" in texto
    assert "Os CONTRATANTES estão cientes" in texto
    # Cada parte carrega o proprio endereco — sem agrupamento de domicilio.
    assert "residentes e domiciliad" not in texto
    assert "residente e domiciliada à Rua Pio Porto" in texto
    assert "com sede à Av. Afonso Pena" in texto


def test_pj_com_homem_usa_plural_masculino(tmp_path):
    gen = ConsumidorGenerator()
    gen.output_dir = tmp_path
    dados = {**UM, "contratantes": [DOIS["contratantes"][1], PJ]}  # Henrique + empresa
    texto = _texto(gen.generate(ContratoConsumidorRequest(**dados))[1])

    assert "doravante denominados apenas CONTRATANTES" in texto


def test_um_contratante_usa_singular(tmp_path):
    gen = ConsumidorGenerator()
    gen.output_dir = tmp_path
    _, path = gen.generate(ContratoConsumidorRequest(**UM))
    texto = _texto(path)

    assert "A CONTRATANTE está ciente e de acordo" in texto
    assert "CONTRATANTES estão cientes" not in texto
    assert "inscrita no CPF sob o n.º 761.794.436-53" in texto
    assert "RG n.º MG-3.764.023" in texto
    assert "residente e domiciliada à Rua Pio Porto" in texto
    # Milheiro: so' a linha da companhia do caso, no valor da faixa da KLM.
    assert "KLM" in texto and "R$ 80,00" in texto
    assert "Gol" not in texto and "Emirates" not in texto
    assert "33.643.420/0001-45" in texto
    # Variante "com reclame aqui": a CONTRATADA elabora a reclamacao.
    assert "elaboração de reclamação para que a CONTRATANTE apresente" in texto
    assert "em até 10 dias" in texto  # prazo padrao


def test_dois_contratantes_usa_plural(tmp_path):
    gen = ConsumidorGenerator()
    gen.output_dir = tmp_path
    _, path = gen.generate(ContratoConsumidorRequest(**DOIS))
    texto = _texto(path)

    assert "Os CONTRATANTES estão cientes e de acordo" in texto
    assert "A CONTRATANTE está ciente e de acordo" not in texto
    # Genero misto -> plural masculino; endereco comum aparece uma vez so'.
    assert "residentes e domiciliados à Rua Major Egido" in texto
    assert texto.count("Rua Major Egido") == 1
    assert "inscrita no CPF sob o n.º 033.395.236-73" in texto
    assert "inscrito no CPF sob o n.º 072.029.566-16" in texto
    assert "doravante denominados apenas CONTRATANTES" in texto
    assert "AZUL" in texto and "R$ 40,00" in texto
    assert "09.296.295/0001-60" in texto
    # Variante sem elaboracao: a CONTRATADA apenas orienta.
    assert "orientação quanto a reclamação extrajudicial a ser apresentada pelos CONTRATANTES" in texto
    assert "em até 5 dias" in texto


def test_nacionalidade_generica_concorda_com_o_genero(tmp_path):
    """O formulario manda "Brasileira (o)"; o contrato tem que escolher um."""
    gen = ConsumidorGenerator()
    gen.output_dir = tmp_path
    dados = {
        **DOIS,
        "contratantes": [
            {**DOIS["contratantes"][0], "nacionalidade": "Brasileira (o)"},
            {**DOIS["contratantes"][1], "nacionalidade": "Brasileira (o)"},
        ],
    }
    texto = _texto(gen.generate(ContratoConsumidorRequest(**dados))[1])

    assert "Brasileira (o)" not in texto
    assert "DANIELA ELIAS ARAUJO MARQUES, brasileira, inscrita" in texto
    assert "HENRIQUE DE OLIVEIRA MARQUES, brasileiro, inscrito" in texto


def test_nacionalidade_estrangeira_e_preservada(tmp_path):
    gen = ConsumidorGenerator()
    gen.output_dir = tmp_path
    dados = {**UM, "contratantes": [{**UM["contratantes"][0], "nacionalidade": "portuguesa"}]}
    assert "portuguesa, inscrita" in _texto(gen.generate(ContratoConsumidorRequest(**dados))[1])


def test_milheiro_pode_ser_ajustado_no_caso(tmp_path):
    """Valor da tabela e' o padrao; o advogado pode sobrescrever no contrato."""
    gen = ConsumidorGenerator()
    gen.output_dir = tmp_path
    _, path = gen.generate(
        ContratoConsumidorRequest(**{**UM, "valor_milheiro_override": 95.5})
    )
    texto = _texto(path)

    assert "R$ 95,50" in texto
    assert "R$ 80,00" not in texto  # o padrao da KLM nao vale mais


def test_varias_companhias_no_polo_passivo(tmp_path):
    """Duas Rés: as duas na Cláusula II e uma linha de milheiro para cada."""
    gen = ConsumidorGenerator()
    gen.output_dir = tmp_path
    dados = {
        **UM,
        "res": [
            {
                "companhia": "KLM",
                "razao_social": "KLM CIA REAL HOLANDESA DE AVIACAO",
                "cnpj": "33.643.420/0001-45",
            },
            {
                "companhia": "Azul",
                "razao_social": "AZUL LINHAS AEREAS BRASILEIRAS S.A.",
                "cnpj": "09.296.295/0001-60",
            },
        ],
    }
    texto = _texto(gen.generate(ContratoConsumidorRequest(**dados))[1])

    assert "em face de KLM CIA REAL HOLANDESA DE AVIACAO" in texto
    assert "e AZUL LINHAS AEREAS BRASILEIRAS S.A." in texto
    assert "33.643.420/0001-45" in texto and "09.296.295/0001-60" in texto
    # Uma linha por companhia, cada uma na sua faixa.
    assert "R$ 80,00" in texto and "R$ 40,00" in texto
    assert "Emirates" not in texto


def test_formato_antigo_de_uma_companhia_ainda_funciona(tmp_path):
    """Contratos gravados antes das múltiplas Rés continuam abrindo."""
    gen = ConsumidorGenerator()
    gen.output_dir = tmp_path
    # UM usa os campos soltos (companhia / re_razao_social / re_cnpj).
    texto = _texto(gen.generate(ContratoConsumidorRequest(**UM))[1])

    assert "KLM CIA REAL HOLANDESA DE AVIACAO" in texto
    assert "R$ 80,00" in texto


def test_formatacao_igual_ao_modelo_do_escritorio(tmp_path):
    """Medido no CONTRATO.docx: Segoe UI 11pt, justificado, 1,15 e margens proprias."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    gen = ConsumidorGenerator()
    gen.output_dir = tmp_path
    doc = Document(gen.generate(ContratoConsumidorRequest(**DOIS))[1])

    secao = doc.sections[0]
    assert round(secao.top_margin.cm, 2) == 0.75
    assert round(secao.bottom_margin.cm, 2) == 1.25
    assert round(secao.left_margin.cm, 2) == 3.00
    assert round(secao.right_margin.cm, 2) == 2.00

    titulo = doc.paragraphs[0]
    assert titulo.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert doc.styles["Heading 1"].font.bold is True

    corpo = next(p for p in doc.paragraphs if p.text.startswith("1 - Os CONTRATANTES"))
    assert corpo.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY
    assert corpo.paragraph_format.line_spacing == 1.15
    assert corpo.runs[0].font.name == "Segoe UI"
    assert corpo.runs[0].font.size.pt == 11

    partes = next(p for p in doc.paragraphs if p.text.startswith("CONTRATANTES:"))
    assert partes.runs[0].bold is True


def test_regeneracao_a_partir_do_json_salvo(tmp_path):
    """Download/edicao/DocuSeal reconstroem o contrato pelo form_data gravado.

    Em FS efemero (serverless) esse e' o caminho normal, nao um caso de borda.
    """
    from app.services.consumidor_generator import ConsumidorGenerator as _CG
    from app.services.contract_dispatch import parse_form_data

    salvo = ContratoConsumidorRequest(**DOIS).model_dump(mode="json")
    request, gen = parse_form_data(salvo)

    assert type(gen) is _CG
    gen.output_dir = tmp_path
    _, path = gen.generate(request, contract_id="reg-1")
    assert "Os CONTRATANTES estão cientes" in _texto(path)


def test_contrato_de_honorarios_nao_muda(tmp_path):
    """O tipo antigo continua roteando para o gerador original."""
    from app.services.contract_dispatch import parse_form_data
    from app.services.contract_generator import ContractGenerator

    _, gen = parse_form_data({"contratantes": [], "escopos": [], "acessorios": {}, "participacao": {}})
    assert type(gen) is ContractGenerator
