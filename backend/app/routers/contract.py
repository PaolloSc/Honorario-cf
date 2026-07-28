from __future__ import annotations

import json
import logging
import re
import uuid
from pathlib import Path

from fastapi import APIRouter, Depends, HTTPException, Query
from fastapi.responses import FileResponse, HTMLResponse
from sqlalchemy.orm import Session

from app.auth import CurrentUser, get_current_user
from app.database import (
    AuditLogDB,
    ContractDB,
    ContractVersionDB,
    ParticipacaoDB,
    get_db,
    serialize_cliente_docs,
    utcnow,
)
from app.models.contract import ContratoRequest, ContratoResponse
from app.services.contract_generator import ContractGenerator


def _infer_tipo_honorario(data: ContratoRequest) -> str:
    """Olha todos escopos. Se >1 tipo distinto → 'misto'; senão usa o único; default 'mensalidade'."""
    tipos = set()
    for esc in data.escopos:
        for h in esc.honorarios:
            v = h.value if hasattr(h, "value") else str(h)
            # Mapeia para vocabulário do calculator
            if v == "hora_trabalhada":
                tipos.add("hora")
            elif v == "pro_labore":
                tipos.add("prolabore")
            elif v == "mensalidade":
                # Verifica subtipo advocacia_partido
                m = esc.mensalidade
                if m and getattr(m, "subtipo", None) and m.subtipo.value == "advocacia_partido":
                    tipos.add("partido")
                else:
                    tipos.add("mensalidade")
            elif v == "exito":
                tipos.add("exito")
            elif v == "permuta":
                tipos.add("exito")  # permuta segue regra de êxito (sem limite)
    if not tipos:
        return "mensalidade"
    if len(tipos) > 1:
        return "misto"
    return next(iter(tipos))


def _extract_cpf_cnpj(data: ContratoRequest) -> str | None:
    if not data.contratantes:
        return None
    first = data.contratantes[0]
    return getattr(first, "cpf", None) or getattr(first, "cnpj", None)


_EMAIL_RE = re.compile(r"[\w\.-]+@[\w\.-]+\.\w+")
_PCT_RE = re.compile(r"(\d+(?:[.,]\d+)?)\s*%")


def _parse_percentual(text: str | None) -> float:
    """Extrai primeiro numero seguido de % do texto livre. '10% de 5000' -> 10.0."""
    if not text:
        return 0.0
    m = _PCT_RE.search(text)
    if not m:
        return 0.0
    try:
        return float(m.group(1).replace(",", "."))
    except ValueError:
        return 0.0


def _parse_email(text: str | None) -> str | None:
    if not text:
        return None
    m = _EMAIL_RE.search(text)
    return m.group(0) if m else None


def _map_natureza_wizard(natureza_text: str | None) -> tuple[bool, bool]:
    """Retorna (eh_captacao, eh_performance) com base no texto livre da wizard."""
    if not natureza_text:
        return False, False
    s = natureza_text.lower()
    eh_cap = "capta" in s
    eh_perf = "perform" in s
    if not eh_cap and not eh_perf:
        # Default: assume captacao se texto livre nao bate
        eh_cap = True
    return eh_cap, eh_perf

logger = logging.getLogger(__name__)
router = APIRouter(prefix="/api/contract", tags=["Contract"])

_generator: ContractGenerator | None = None


def get_generator() -> ContractGenerator:
    global _generator
    if _generator is None:
        _generator = ContractGenerator()
    return _generator


def _extract_client_info(data: ContratoRequest) -> tuple[str, str]:
    if not data.contratantes:
        return ("", "")
    first = data.contratantes[0]
    name = getattr(first, "nome", None) or getattr(first, "razao_social", "")
    email = getattr(first, "email", "")
    return (name, email)


@router.post("/generate", response_model=ContratoResponse)
def generate_contract(
    data: ContratoRequest,
    user: CurrentUser = Depends(get_current_user),
    db: Session = Depends(get_db),
) -> ContratoResponse:
    try:
        gen = get_generator()
        contract_id, filepath = gen.generate(data)

        client_name, client_email = _extract_client_info(data)
        form_dict = data.model_dump(mode="json")

        contract = ContractDB(
            contract_id=contract_id,
            status="rascunho",
            client_name=client_name,
            client_email=client_email,
            current_version=1,
            created_by=user.email,
            updated_by=user.email,
            created_at=utcnow(),
            updated_at=utcnow(),
            cliente_docs=serialize_cliente_docs(form_dict),
        )
        db.add(contract)

        form_data_json = json.dumps(form_dict, ensure_ascii=False)
        version = ContractVersionDB(
            contract_id=contract_id,
            version_number=1,
            form_data_json=form_data_json,
            file_path=filepath,
            created_by=user.email,
            created_at=utcnow(),
        )
        db.add(version)

        audit = AuditLogDB(
            contract_id=contract_id,
            action="criacao",
            detail=f"Contrato gerado para {client_name}",
            version_number=1,
            user_email=user.email,
            created_at=utcnow(),
        )
        db.add(audit)

        # ── Plano A: cria Participação rascunho automática para o financeiro ──
        # Percentuais zerados + aprovada=False. Setor financeiro complementa e aprova.
        try:
            tipo = _infer_tipo_honorario(data)
            participacao_data = data.participacao

            # Parse valores da wizard
            pct = _parse_percentual(
                participacao_data.percentual_ou_valor if participacao_data else None
            )
            eh_cap, eh_perf = _map_natureza_wizard(
                participacao_data.natureza if participacao_data else None
            )
            pct_captacao = pct if eh_cap else 0.0
            pct_performance = pct if eh_perf else 0.0

            # Email: prioriza campo estruturado novo, fallback parse do legado, fallback user.
            email_novo = (participacao_data.contato_financeiro_email or "").strip() if participacao_data else ""
            email_wizard = email_novo or _parse_email(
                participacao_data.contato_financeiro_cliente if participacao_data else None
            )
            beneficiario_email = email_wizard or user.email
            # Nome: para_quem agora e' lista; junta nomes. Fallback user.name.
            para_quem_nomes = (participacao_data.para_quem if participacao_data else None) or []
            beneficiario_nome = ", ".join(para_quem_nomes) if para_quem_nomes else user.name

            # Backend natureza = contratual/societario (campo legal, nao wizard)
            natureza_val = "contratual"

            # Motivos: usa texto da wizard quando aplicável
            motivo_cap = (
                participacao_data.responsavel_captacao
                if participacao_data and pct_captacao > 0 and participacao_data.responsavel_captacao
                else None
            )
            motivo_perf = (
                participacao_data.responsavel_gestao
                if participacao_data and pct_performance > 0 and participacao_data.responsavel_gestao
                else None
            )

            obs_extra = ""
            if participacao_data:
                if participacao_data.valor_tipo == "percentual" and participacao_data.valor_percentual:
                    valor_str = f"{participacao_data.valor_percentual}%"
                elif participacao_data.valor_tipo == "valor" and participacao_data.valor_monetario is not None:
                    valor_str = f"R$ {participacao_data.valor_monetario:.2f}"
                elif participacao_data.valor_tipo == "outro" and participacao_data.valor_outro:
                    valor_str = participacao_data.valor_outro
                else:
                    valor_str = participacao_data.percentual_ou_valor or "-"
                obs_extra = (
                    f"Rascunho automatico do wizard. "
                    f"Captacao responsavel: {participacao_data.responsavel_captacao or '-'} · "
                    f"Gestao: {participacao_data.responsavel_gestao or '-'} · "
                    f"Valor wizard: {valor_str} · "
                    f"Natureza wizard: {participacao_data.natureza or '-'}"
                )
            else:
                obs_extra = "Rascunho automatico do wizard."

            rascunho = ParticipacaoDB(
                contract_id=contract_id,
                beneficiario_email=beneficiario_email,
                beneficiario_nome=beneficiario_nome or "",
                tipo_honorario=tipo,
                percentual_captacao=pct_captacao,
                percentual_performance=pct_performance,
                motivo_captacao=motivo_cap,
                motivo_performance=motivo_perf,
                natureza=natureza_val,
                cliente_cpf_cnpj=_extract_cpf_cnpj(data),
                data_inicio=utcnow().date(),
                vinculo_ativo=True,
                aprovada=False,
                observacoes=obs_extra,
                created_by=user.email,
                created_at=utcnow(),
            )
            db.add(rascunho)
            db.add(
                AuditLogDB(
                    contract_id=contract_id,
                    action="participacao_rascunho",
                    detail=f"Rascunho automático criado para revisão do financeiro (tipo={tipo})",
                    user_email=user.email,
                    created_at=utcnow(),
                )
            )
        except Exception as part_err:
            # Não bloqueia geração de contrato se falhar criação do rascunho
            logger.warning("Falha ao criar participação rascunho automática: %s", part_err)

        db.commit()

        logger.info("Contract generated by %s: %s", user.email, contract_id)

        return ContratoResponse(
            success=True,
            message="Contrato gerado com sucesso",
            contract_id=contract_id,
            download_url=f"/api/contract/{contract_id}/download",
        )
    except Exception as e:
        db.rollback()
        logger.error("Failed to generate contract: %s", e)
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/{contract_id}/download")
def download_contract(
    contract_id: str,
    version: int | None = None,
    user: CurrentUser = Depends(get_current_user),
    db: Session = Depends(get_db),
) -> FileResponse:
    try:
        uuid.UUID(contract_id)
    except ValueError:
        raise HTTPException(400, "ID de contrato inválido")

    # Check ownership or admin
    contract = db.query(ContractDB).filter(ContractDB.contract_id == contract_id).first()
    if not contract:
        raise HTTPException(404, "Contrato nao encontrado")
    if user.role != "admin" and contract.created_by != user.email:
        raise HTTPException(403, "Sem permissao")

    # Find the file path from DB (specific version or latest)
    query = db.query(ContractVersionDB).filter(ContractVersionDB.contract_id == contract_id)
    if version:
        ver = query.filter(ContractVersionDB.version_number == version).first()
    else:
        ver = query.order_by(ContractVersionDB.version_number.desc()).first()

    gen = get_generator()
    output_dir = Path(gen.output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    filepath: Path | None = None

    if ver and ver.file_path:
        stored = Path(ver.file_path)
        if stored.exists():
            filepath = stored
        else:
            candidate = output_dir / stored.name
            if candidate.exists():
                filepath = candidate

    if filepath is None:
        fallback = output_dir / f"contrato_{contract_id}.docx"
        if fallback.exists():
            filepath = fallback

    # Strategy 4: Regenerate from form_data_json (handles ephemeral filesystems)
    if filepath is None and ver and ver.form_data_json:
        logger.warning(
            "Contract file not on disk for %s. Regenerating from stored form data...",
            contract_id,
        )
        try:
            form_data = json.loads(ver.form_data_json)
            contrato_data = ContratoRequest(**form_data)
            _, new_filepath = gen.generate(contrato_data, contract_id=contract_id)
            filepath = Path(new_filepath)
            # Update stored path
            ver.file_path = str(filepath)
            db.commit()
            logger.info("Regenerated contract file at: %s", filepath)
        except FileNotFoundError as template_err:
            # Template not found (ephemeral FS) - create a minimal placeholder DOCX
            logger.warning("Template not found, creating minimal DOCX: %s", template_err)
            try:
                from docx import Document as _Doc
                doc = _Doc()
                doc.add_paragraph("Contrato em processamento - documento sera regenerado.")
                minimal_path = output_dir / f"contrato_{contract_id}.docx"
                doc.save(str(minimal_path))
                ver.file_path = str(minimal_path)
                db.commit()
                filepath = minimal_path
                logger.info("Created minimal placeholder DOCX at: %s", minimal_path)
            except Exception as min_err:
                logger.error("Failed to create minimal DOCX: %s", min_err)
        except Exception as regen_err:
            logger.error("Failed to regenerate contract %s: %s", contract_id, regen_err)

    if filepath is None or not filepath.exists():
        logger.error(
            "Contract file not found for %s. Tried: stored=%s, output_dir=%s",
            contract_id,
            ver.file_path if ver else "N/A",
            output_dir,
        )
        raise HTTPException(status_code=404, detail="Arquivo nao encontrado")

    return FileResponse(
        path=str(filepath),
        filename=_contract_filename(contract.client_name, contract_id),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


import re as _re

# DocuSeal merge-tag (ex.: "{{Assinatura Fulano;type=signature;role=Contratante}}").
# No preview trocamos por uma linha de assinatura para nao vazar a sintaxe crua.
_SIG_TAG = _re.compile(r"\{\{[^}]*?type=signature[^}]*?\}\}")

_FILENAME_BAD = _re.compile(r'[\r\n\t/\\:*?"<>|]+')


def _safe_filename_part(name: str) -> str:
    """Sanitiza um trecho vindo do usuário para uso em nome de arquivo/anexo."""
    cleaned = " ".join(_FILENAME_BAD.sub(" ", name).split())
    return cleaned[:120].strip()


def _contract_filename(client_name: str | None, contract_id: str) -> str:
    nome = _safe_filename_part(client_name) if client_name else ""
    return f"Contrato Honorários — {nome}.docx" if nome else f"contrato_honorarios_{contract_id}.docx"


def _clean_preview_text(text: str) -> str:
    # A tag do DocuSeal e' invisivel no Word (texto branco) e o documento ja traz
    # a linha de assinatura; deixa-la virar underscores duplicava a linha na previa.
    return _SIG_TAG.sub("", text)


def _clause_level(paragraph_el) -> int | None:
    """Nível do parágrafo na lista de cláusulas, ou None se não for numerado."""
    from docx.oxml.ns import qn

    num_pr = paragraph_el.find(qn("w:pPr") + "/" + qn("w:numPr"))
    if num_pr is None:
        return None
    ilvl = num_pr.find(qn("w:ilvl"))
    return int(ilvl.get(qn("w:val"))) if ilvl is not None else 0


def _docx_to_html(filepath: Path) -> str:
    """Render the generated DOCX as simple HTML for inline preview (no external deps)."""
    from html import escape

    from docx import Document
    from docx.oxml.ns import qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    doc = Document(str(filepath))
    parts: list[str] = []
    contadores = [0, 0, 0]  # niveis da lista de clausulas (o Word numera no .docx)
    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            p = Paragraph(child, doc)
            text = escape(_clean_preview_text(p.text))
            if not text.strip():
                continue
            ilvl = _clause_level(child)
            if ilvl is not None:
                contadores[ilvl] += 1
                for abaixo in range(ilvl + 1, len(contadores)):
                    contadores[abaixo] = 0
                text = ".".join(str(n) for n in contadores[: ilvl + 1]) + ". " + text
            style = p.style.name if p.style else ""
            if style == "Heading 1":
                parts.append(f"<h1>{text}</h1>")
            elif style.startswith("Heading"):
                parts.append(f"<h3>{text}</h3>")
            else:
                parts.append(f"<p>{text}</p>")
        elif child.tag == qn("w:tbl"):
            t = Table(child, doc)
            has_borders = t._tbl.tblPr.first_child_found_in("w:tblBorders") is not None
            # Tabelas com borda tem linha de titulo (Escopo/Preco etc.) -> th centralizado.
            def celula_html(cell) -> str:
                # Cada paragrafo da celula e' uma linha. `cell.text` junta tudo num
                # texto so e o HTML colapsa a quebra: o rotulo da assinatura acabava
                # na mesma linha dos underscores.
                linhas = [
                    escape(_clean_preview_text(p.text)).strip() for p in cell.paragraphs
                ]
                return "<br>".join(linha for linha in linhas if linha)

            rows_html = "".join(
                "<tr>" + "".join(
                    f"<{'th' if has_borders and i == 0 else 'td'}>{celula_html(c)}"
                    f"</{'th' if has_borders and i == 0 else 'td'}>"
                    for c in row.cells
                ) + "</tr>"
                for i, row in enumerate(t.rows)
            )
            css = "" if has_borders else ' class="noborder"'
            parts.append(f"<table{css}>{rows_html}</table>")
    return (
        "<!doctype html><html><head><meta charset='utf-8'>"
        "<style>"
        "body{font-family:'Segoe UI',sans-serif;max-width:800px;margin:2rem auto;"
        "line-height:1.5;color:#000;background:#fff;padding:0 1rem}"
        "h1{text-align:center;font-size:1.1rem}"
        "h3{font-size:1rem;margin-top:1.2rem}"
        "table{border-collapse:collapse;width:100%;margin:.5rem 0}"
        "td,th{border:1px solid #000;padding:6px;font-size:.9rem}"
        "th{text-align:center}"
        "table.noborder td{border:none}"
        "</style></head>"
        f"<body>{''.join(parts)}</body></html>"
    )


@router.get("/{contract_id}/preview", response_class=HTMLResponse)
def preview_contract(
    contract_id: str,
    user: CurrentUser = Depends(get_current_user),
    db: Session = Depends(get_db),
) -> HTMLResponse:
    """Inline HTML preview of the latest contract version."""
    try:
        uuid.UUID(contract_id)
    except ValueError:
        raise HTTPException(400, "ID de contrato inválido")

    contract = db.query(ContractDB).filter(ContractDB.contract_id == contract_id).first()
    if not contract:
        raise HTTPException(404, "Contrato nao encontrado")
    if user.role != "admin" and contract.created_by != user.email:
        raise HTTPException(403, "Sem permissao")

    from app.routers.docuseal import _resolve_contract_filepath

    filepath = _resolve_contract_filepath(contract_id, db)
    return HTMLResponse(_docx_to_html(Path(filepath)))
